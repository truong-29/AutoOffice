import os
import docx
from docx.enum.section import WD_SECTION_START
import datetime
import traceback
import shutil
import win32com.client
import sys
import platform
import gc
import pythoncom
import tempfile
import time
import logging

class WordProcessor:
    # Biến lưu thể hiện đã được tạo (singleton pattern)
    _instance = None
    
    def __new__(cls, *args, **kwargs):
        """
        Đảm bảo chỉ có một thể hiện của WordProcessor được tạo ra
        """
        if cls._instance is None:
            cls._instance = super(WordProcessor, cls).__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self, log_file="word_processor_log.txt"):
        """
        Khởi tạo đối tượng xử lý file Word
        
        Args:
            log_file (str): Đường dẫn đến file log
        """
        # Không khởi tạo lại nếu đã được khởi tạo
        if getattr(self, '_initialized', False):
            self.log(f"Tái sử dụng thể hiện WordProcessor đã có - {id(self)}")
            return
            
        self._initialized = True
        self.log_file = log_file
        self._word_apps = []  # Danh sách các thể hiện Word COM đã tạo
        
        # Ghi log khởi động
        self.log("Đang khởi tạo WordProcessor")
        self.log(f"Phiên bản Python: {sys.version}")
        self.log(f"Đường dẫn Python: {sys.executable}")
        
        # Ghi thông tin hệ thống
        try:
            self.log(f"Hệ điều hành: {platform.platform()}")
            self.log(f"Tên máy: {platform.node()}")
            self.log(f"Mã hóa hệ thống: {sys.getfilesystemencoding()}")
        except:
            self.log("Không thể lấy thông tin hệ thống chi tiết")
            
        self.log(f"Đã khởi tạo WordProcessor với ID: {id(self)}")
        
    def log(self, message, error=False):
        """
        Ghi thông tin vào file log
        
        Args:
            message (str): Thông điệp cần ghi
            error (bool): True nếu là lỗi, mặc định là False
        """
        # Tạo thư mục chứa file log nếu chưa tồn tại
        log_dir = os.path.dirname(self.log_file)
        if log_dir and not os.path.exists(log_dir):
            try:
                os.makedirs(log_dir)
            except Exception as e:
                print(f"Không thể tạo thư mục log: {e}")
        
        # Xác định loại thông điệp
        log_type = "ERROR" if error else "INFO"
        
        # Lấy thời gian hiện tại
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Chuẩn bị nội dung log
        log_entry = f"[{timestamp}] [{log_type}] {message}"
        
        # In ra console để debug
        print(log_entry)
        
        # Thêm stack trace nếu là lỗi
        if error:
            stack_trace = traceback.format_stack()
            stack_info = "".join(stack_trace[:-1])  # Loại bỏ frame hiện tại
            log_entry += f"\nStack trace:\n{stack_info}"
        
        # Thử ghi vào file log
        try:
            with open(self.log_file, "a", encoding="utf-8") as f:
                f.write(log_entry + "\n")
        except Exception as e:
            # Nếu không thể ghi vào file log, in ra console
            print(f"Không thể ghi vào file log: {e}")
            # Thử tạo file log mới ở thư mục hiện tại
            try:
                fallback_log = "word_processor_error.log"
                with open(fallback_log, "a", encoding="utf-8") as f:
                    f.write(log_entry + "\n")
                    f.write(f"Lỗi khi ghi vào file log gốc: {e}\n")
            except:
                pass
        
    def open_document(self, file_path):
        """
        Mở tệp Word và trả về True nếu mở thành công
        
        Args:
            file_path (str): Đường dẫn đến tệp Word
            
        Returns:
            bool: True nếu mở thành công, False nếu có lỗi
        """
        self.log(f"Đang mở tệp: {file_path}")
        try:
            self.file_path = file_path
            self.document = docx.Document(file_path)
            self.log(f"Đã mở tệp thành công: {file_path}")
            return True
        except Exception as e:
            error_msg = f"Lỗi khi mở tệp: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            return False
    
    def analyze_document(self):
        """
        Phân tích tài liệu tìm kiếm section breaks
        
        Returns:
            list: Danh sách các section và thông tin của chúng
        """
        self.log("Bắt đầu phân tích tài liệu")
        if not self.document:
            self.log("Không có tài liệu nào được mở", error=True)
            return []
        
        results = []
        try:
            # Phân tích số lượng section và loại section break
            self.log(f"Số section trong tài liệu: {len(self.document.sections)}")
            for i, section in enumerate(self.document.sections):
                section_type = "Continuous"
                
                # Xác định loại section break
                if section.start_type == WD_SECTION_START.NEW_PAGE:
                    section_type = "Next Page"
                elif section.start_type == WD_SECTION_START.NEW_COLUMN:
                    section_type = "New Column"
                elif section.start_type == WD_SECTION_START.EVEN_PAGE:
                    section_type = "Even Page"
                elif section.start_type == WD_SECTION_START.ODD_PAGE:
                    section_type = "Odd Page"
                
                likely_blank = section_type in ["Next Page", "Even Page", "Odd Page"]
                self.log(f"Section {i+1}: Kiểu={section_type}, Có thể gây trang trắng={likely_blank}")
                    
                results.append({
                    "index": i + 1,
                    "type": section_type,
                    "likely_blank": likely_blank
                })
            
            self.log(f"Đã phân tích xong: tìm thấy {len(results)} section")
            return results
        except Exception as e:
            error_msg = f"Lỗi khi phân tích tài liệu: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            return []
    
    def detect_blank_pages(self, sections_results):
        """
        Phát hiện chính xác các trang trắng trong tài liệu
        
        Args:
            sections_results (list): Kết quả phân tích section ban đầu
            
        Returns:
            tuple: (Kết quả phân tích section đã cập nhật, Danh sách chi tiết từng trang)
        """
        self.log("Đang phát hiện trang trắng chi tiết...")
        
        if not self.document or not sections_results:
            self.log("Không có tài liệu hoặc kết quả section để phân tích", error=True)
            return sections_results, []
            
        word_app = None
        temp_file = None
        all_pages_info = []  # Khởi tạo danh sách chứa thông tin chi tiết từng trang
        
        # Đảm bảo danh sách lưu trữ file tạm tồn tại
        if not hasattr(self, 'temp_files'):
            self.temp_files = []
        
        try:
            # Sử dụng Word COM để kiểm tra trang trắng
            import win32com.client
            import pythoncom
            
            # Khởi tạo COM cho thread hiện tại nếu cần
            try:
                pythoncom.CoInitialize()
            except:
                pass
            
            # Lưu tệp tạm thời
            temp_file = self.save_temp_file()
            if not temp_file or not os.path.exists(temp_file):
                self.log("Không thể tạo file tạm để phân tích", error=True)
                return sections_results, []
                
            # Lưu để xóa sau
            self.temp_files.append(temp_file)
            
            # Tạo instance Word mới
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            self._word_apps.append(word_app)  # Lưu lại để có thể đóng sau
            
            # Mở file tạm
            doc = word_app.Documents.Open(temp_file)
                
            # Lấy số trang trong tài liệu
            total_pages = doc.ComputeStatistics(2)  # wdStatisticPages = 2
            self.log(f"Tổng số trang trong tài liệu: {total_pages}")
                
            # Duyệt qua từng trang để phân tích
            active_window = doc.ActiveWindow
            active_window.View.Type = 3  # wdPrintView = 3
            
            # Danh sách thông tin chi tiết các trang
            all_pages_info = []
            
            # Duyệt qua từng trang để kiểm tra
            for page_num in range(1, total_pages + 1):
                # Di chuyển đến trang cần kiểm tra
                active_window.ActivePane.View.GoTo(1, page_num, 1, None)  # wdGoToPage = 1
                
                # Thông tin chi tiết về trang
                page_info = {
                    "page_number": page_num,
                    "is_blank": False,
                    "content_type": "unknown",
                    "has_text": False,
                    "has_table": False,
                    "has_image": False,
                    "has_footer": False,
                    "has_header": False,
                    "has_section_break": False,
                    "is_manual_delete": False,
                    "section_index": -1
                }
                
                # Lấy tất cả nội dung trên trang
                try:
                    word_app.Selection.GoTo(1, page_num, 1, None)  # wdGoToPage = 1
                    word_app.Selection.Expand(0)  # wdStory = 0, mở rộng selection đến toàn bộ trang
                    
                    text = word_app.Selection.Text
                    tables_count = word_app.Selection.Tables.Count
                    shapes_count = word_app.Selection.ShapeRange.Count
                    
                    # Kiểm tra có chữ không
                    page_info["has_text"] = len(text.strip()) > 0
                    
                    # Kiểm tra có bảng không
                    page_info["has_table"] = tables_count > 0
                    
                    # Kiểm tra có hình ảnh không
                    page_info["has_image"] = shapes_count > 0
                    
                    # Xác định kiểu nội dung
                    if page_info["has_text"] or page_info["has_table"] or page_info["has_image"]:
                        page_info["content_type"] = "content"
                    else:
                        page_info["content_type"] = "blank"
                        page_info["is_blank"] = True
                except Exception as e:
                    self.log(f"Lỗi khi phân tích nội dung trang {page_num}: {e}", error=True)
                            
                # Tìm section chứa trang hiện tại
                for i, section in enumerate(sections_results):
                    if page_num >= section["start_page"] and page_num <= section["end_page"]:
                        page_info["section_index"] = i
                        break
                
                # Thêm vào danh sách kết quả
                all_pages_info.append(page_info)
            
            # Phân tích thêm để xác định trang trắng đặc biệt (ví dụ: trang cuối section)
            for page_idx, page_info in enumerate(all_pages_info):
                page_num = page_info["page_number"]
                
                # Tìm trang trước và sau
                prev_page = next((p for p in all_pages_info if p["page_number"] == page_num - 1), None)
                next_page = next((p for p in all_pages_info if p["page_number"] == page_num + 1), None)
                
                # Kiểm tra xem trang hiện tại có phải là trang cuối của section không
                if page_info["section_index"] >= 0:
                    section = sections_results[page_info["section_index"]]
                    if page_num == section["end_page"] and next_page and next_page["section_index"] != page_info["section_index"]:
                        page_info["has_section_break"] = True
                        
                        # Nếu là trang trắng ở cuối section, đánh dấu để xóa thủ công sau
                        if page_info["is_blank"]:
                            page_info["is_manual_delete"] = True
                            self.log(f"Trang {page_num} là trang trắng cuối section, cần xóa thủ công")
                
            # Cập nhật thông tin sections_results
            for section_idx, section in enumerate(sections_results):
                section_blank_pages = []
                    
                # Tìm các trang trắng trong section này
                for page in all_pages_info:
                    if page["section_index"] == section_idx and page["is_blank"]:
                        section_blank_pages.append(page["page_number"])
                
                # Cập nhật thông tin
                section["blank_pages"] = section_blank_pages
                section["blank_page_count"] = len(section_blank_pages)
            
            # Đóng tài liệu sau khi phân tích xong
            doc.Close(False)  # False = không lưu thay đổi
            
            return sections_results, all_pages_info
        except Exception as e:
            self.log(f"Lỗi khi phát hiện trang trắng: {e}", error=True)
            self.log(traceback.format_exc(), error=True)
            return sections_results, []
        finally:
            # Đóng Word nếu đã mở
            if word_app:
                try:
                    word_app.Quit()
                except:
                    pass
    
    def fix_document(self, sections_to_fix=None):
        """
        Sửa tài liệu bằng cách chuyển các section breaks thành Continuous
        
        Args:
            sections_to_fix (list, optional): Danh sách các section cần sửa.
                                           Nếu None, sửa tất cả các section.
        
        Returns:
            str/bool: Đường dẫn đến tệp đã sửa nếu thành công, False nếu thất bại
        """
        if not self.document:
            self.log("Không có tài liệu nào được mở", error=True)
            return False
        
        # Phân tích document để lấy thông tin section và trang trắng
        sections_results = self.analyze_document()
        if not sections_results:
            self.log("Không thể phân tích tài liệu", error=True)
            return False
            
        # Phát hiện trang trắng chi tiết
        updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
        
        # Đảm bảo sections_to_fix là list số nguyên
        if sections_to_fix is not None:
            try:
                # Chuyển đổi tất cả các phần tử thành số nguyên
                sections_to_fix = [int(s) for s in sections_to_fix]
                section_info = f"các section: {sections_to_fix}"
            except (ValueError, TypeError) as e:
                self.log(f"Lỗi chuyển đổi sections_to_fix thành số nguyên: {e}", error=True)
                return False
        else:
            # Nếu không chỉ định section cụ thể, phân tích để xác định section cần sửa
            sections_with_blank_pages = []
            blank_page_sections = set()
            
            # Lấy tất cả các section có trang trắng
            for section in updated_sections:
                if section.get("blank_pages"):
                    sections_with_blank_pages.append(section)
                    blank_page_sections.add(int(section["index"]))
            
            # Lấy thêm các section có ảnh hưởng đến trang trắng (kể cả không phải section chứa trang trắng)
            for page_info in all_pages_info:
                if page_info.get("is_blank") and page_info.get("section_number"):
                    # Nếu trang trắng có section break, kiểm tra section tiếp theo
                    if page_info.get("contains_section_break") and page_info.get("section_number") < len(updated_sections):
                        next_section = page_info.get("section_number") + 1
                        blank_page_sections.add(next_section)
            
            # Thêm các section tiếp theo sau section có trang trắng (có thể ảnh hưởng)
            affected_sections = set()
            for section_idx in blank_page_sections:
                affected_sections.add(section_idx)
                # Thêm section tiếp theo nếu có
                if section_idx < len(updated_sections):
                    affected_sections.add(section_idx + 1)
            
            # Kết hợp tất cả các section cần xử lý và lọc chỉ lấy những section > 1
            sections_to_fix = [idx for idx in affected_sections if idx > 1]
            
            # Sắp xếp theo thứ tự tăng dần
            sections_to_fix.sort()
            
            section_info = f"các section ảnh hưởng đến trang trắng: {sections_to_fix}"
        
        # Lọc tiếp để chỉ lấy section không phải continuous
        sections_to_fix_filtered = []
        for section_idx in sections_to_fix:
            if section_idx > 1 and section_idx <= len(updated_sections):
                section = updated_sections[section_idx - 1]  # -1 vì index bắt đầu từ 0 nhưng section_idx từ 1
                if section["type"] != "Continuous":
                    sections_to_fix_filtered.append(section_idx)
        
        if not sections_to_fix_filtered:
            self.log("Không có section nào cần sửa (tất cả đã là Continuous hoặc không ảnh hưởng đến trang trắng).")
            return False
            
        sections_to_fix = sections_to_fix_filtered
        self.log(f"Bắt đầu sửa {section_info}")
        
        try:
            # Lưu tệp tạm thời để xử lý bằng phương pháp khác
            temp_file = self.file_path + ".temp.docx"
            self.document.save(temp_file)
            self.log(f"Đã lưu tệp tạm thời: {temp_file}")
            
            # Mở tệp Word bằng cách sử dụng API COM (chỉ hoạt động trên Windows)
            self.log("Đang khởi tạo Word COM Application")
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            try:
                self.log(f"Đang mở tệp bằng Word COM: {temp_file}")
                doc = word_app.Documents.Open(os.path.abspath(temp_file))
                
                # Ghi log thông tin về sections
                self.log(f"Số section trong tài liệu (COM): {doc.Sections.Count}")
                
                sections_fixed = 0
                # Duyệt qua các section trong tài liệu
                for i in range(2, doc.Sections.Count + 1):
                    # Nếu sections_to_fix không được chỉ định, sửa tất cả
                    # Hoặc nếu section này nằm trong danh sách cần sửa
                    if sections_to_fix is None or (int(i) in sections_to_fix):
                        current_type = doc.Sections(i).PageSetup.SectionStart
                        self.log(f"Đang sửa section {i}: Loại hiện tại={current_type} -> Continuous (0)")
                        # Đặt kiểu section break thành Continuous (0)
                        doc.Sections(i).PageSetup.SectionStart = 0
                        sections_fixed += 1
                
                # Lưu và đóng tài liệu
                self.log(f"Đang lưu tài liệu đã sửa (đã sửa {sections_fixed} section)")
                doc.Save()
                doc.Close()
                
                # Tải lại tài liệu đã sửa
                self.log(f"Đang tải lại tài liệu từ: {temp_file}")
                self.document = docx.Document(temp_file)
                
                # Đổi tên tệp tạm thành tệp gốc (hoặc tạo một tệp mới)
                file_name = os.path.basename(self.file_path)
                file_dir = os.path.dirname(self.file_path)
                name_without_ext = os.path.splitext(file_name)[0]
                output_path = os.path.join(file_dir, f"{name_without_ext}_fixed.docx")
                
                self.log(f"Đang lưu tệp đã sửa vào: {output_path}")
                self.document.save(output_path)
                
                # Xóa tệp tạm
                self.log(f"Đang xóa tệp tạm: {temp_file}")
                os.remove(temp_file)
                
                self.log(f"Đã hoàn tất sửa tài liệu: {output_path}")
                return output_path
                
            finally:
                self.log("Đang đóng Word COM Application")
                word_app.Quit()
                
        except Exception as e:
            error_msg = f"Lỗi khi sửa tài liệu: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            return False

    def fix_sections(self, file_path, section_indexes=None, output_path=None):
        """
        Sửa lỗi section break trong tài liệu Word
        Chuyển các section break thành Continuous
        
        Args:
            file_path (str): Đường dẫn đến file Word
            section_indexes (list, optional): Danh sách các section cần sửa. Nếu None, sửa tất cả.
            output_path (str, optional): Đường dẫn lưu file kết quả. Nếu None, tạo tên tự động.
            
        Returns:
            str: Đường dẫn đến file đã được sửa
        """
        try:
            # Đảm bảo section_indexes là list số nguyên nếu được cung cấp
            if section_indexes is not None:
                try:
                    # Chuyển đổi tất cả các phần tử thành số nguyên
                    section_indexes = [int(s) for s in section_indexes]
                except (ValueError, TypeError) as e:
                    self.log(f"Lỗi chuyển đổi section_indexes thành số nguyên: {e}", error=True)
                    return False
            
            # Mở file Word
            self.open_document(file_path)
            if not self.document:
                self.log("Không thể mở tài liệu Word", error=True)
                return False
                
            # Phân tích tài liệu để tìm các section
            sections_results = self.analyze_document()
            if not sections_results:
                self.log("Không thể phân tích tài liệu", error=True)
                return False
                
            # Phát hiện trang trắng chi tiết
            updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
            
            # Lọc các section cần sửa dựa trên section_indexes đã chỉ định
            if section_indexes:
                sections_to_fix = [s for s in updated_sections if int(s["index"]) in section_indexes]
            else:
                # Nếu không chỉ định section cụ thể, lọc tất cả section có thể gây trang trắng
                # và cả những section có trang trắng (có thể khác section gây ra trang trắng)
                sections_with_blank_pages = []
                blank_page_sections = set()
                
                # Lấy tất cả các section có trang trắng
                for section in updated_sections:
                    if section.get("blank_pages"):
                        sections_with_blank_pages.append(section)
                        blank_page_sections.add(int(section["index"]))
                
                # Lấy thêm các section có ảnh hưởng đến trang trắng (kể cả không phải section chứa trang trắng)
                for page_info in all_pages_info:
                    if page_info.get("is_blank") and page_info.get("section_number"):
                        # Nếu trang trắng có section break, kiểm tra section tiếp theo
                        if page_info.get("contains_section_break") and page_info.get("section_number") < len(updated_sections):
                            next_section = page_info.get("section_number") + 1
                            blank_page_sections.add(next_section)
                
                # Thêm các section tiếp theo sau section có trang trắng (có thể ảnh hưởng)
                affected_sections = set()
                for section_idx in blank_page_sections:
                    affected_sections.add(section_idx)
                    # Thêm section tiếp theo nếu có
                    if section_idx < len(updated_sections):
                        affected_sections.add(section_idx + 1)
                
                # Thêm các section có thể gây ra trang trắng dựa trên kiểu section break
                likely_blank_sections = [s for s in updated_sections if s.get("likely_blank", False)]
                
                # Kết hợp tất cả các section cần xử lý
                all_sections_to_check = []
                all_sections_to_check.extend(sections_with_blank_pages)
                all_sections_to_check.extend([s for s in updated_sections if int(s["index"]) in affected_sections])
                all_sections_to_check.extend(likely_blank_sections)
                
                # Loại bỏ trùng lặp và giữ thứ tự
                seen = set()
                sections_to_fix = []
                for section in all_sections_to_check:
                    if int(section["index"]) not in seen:
                        seen.add(int(section["index"]))
                        sections_to_fix.append(section)
            
            # Lọc thêm: chỉ lấy các section không phải Continuous và không phải section đầu tiên
            sections_to_fix = [s for s in sections_to_fix if s["type"] != "Continuous" and int(s["index"]) > 1]
            
            # Ghi log chi tiết
            self.log(f"Tổng số section được phát hiện: {len(updated_sections)}")
            self.log(f"Tổng số trang trắng được phát hiện: {len([p for p in all_pages_info if p.get('is_blank', False)])}")
            
            if not sections_to_fix:
                self.log("Không có section nào cần sửa (tất cả đã là Continuous hoặc không được chọn).")
                return file_path
            
            self.log(f"Số section cần sửa: {len(sections_to_fix)}")
            for section in sections_to_fix:
                self.log(f"- Section {section['index']}: {section['type']} -> Continuous")
                
            # Tạo tên file kết quả
            if not output_path:
                file_dir = os.path.dirname(file_path)
                file_name = os.path.basename(file_path)
                base_name, ext = os.path.splitext(file_name)
                output_path = os.path.join(file_dir, f"{base_name}_fixed{ext}")
                
            # Mở tài liệu gốc để sửa
            doc = docx.Document(file_path)
            
            # Duyệt qua các section trong tài liệu
            sections_changed = 0
            for i, section in enumerate(doc.sections):
                section_idx = i + 1
                # Kiểm tra xem section này có cần sửa không
                if section_idx > 1 and section_idx in [int(s["index"]) for s in sections_to_fix]:
                    current_type = section.start_type
                    section_name = self._get_section_type_name(current_type)
                    self.log(f"Đang sửa section {section_idx}: {section_name} -> Continuous")
                    # Đặt kiểu section break thành Continuous (0)
                    section.start_type = WD_SECTION_START.CONTINUOUS
                    sections_changed += 1
            
            # Lưu tài liệu đã sửa vào file mới
            self.log(f"Đã sửa {sections_changed} section(s)")
            self.log(f"Đang lưu tài liệu đã sửa vào: {output_path}")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            error_msg = f"Lỗi khi sửa tài liệu: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            return False

    def remove_blank_pages(self, file_path, pages_to_remove=None, output_path=None):
        """
        Xóa trực tiếp các trang trắng từ tài liệu Word
        
        Args:
            file_path (str): Đường dẫn đến file Word
            pages_to_remove (list, optional): Danh sách các trang cần xóa. Nếu None, xóa tất cả trang trắng.
            output_path (str, optional): Đường dẫn lưu file kết quả. Nếu None, tạo tên tự động.
            
        Returns:
            str: Đường dẫn đến file đã được sửa
        """
        try:
            # Mở file Word
            self.open_document(file_path)
            if not self.document:
                self.log("Không thể mở tài liệu Word", error=True)
                return False
                
            # Phân tích tài liệu để tìm các section
            sections_results = self.analyze_document()
            if not sections_results:
                self.log("Không thể phân tích tài liệu", error=True)
                return False
                
            # Phát hiện trang trắng chi tiết
            updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
            
            # Nếu không chỉ định cụ thể trang cần xóa, lấy tất cả trang trắng
            if pages_to_remove is None:
                # Lấy số trang của tất cả trang được đánh dấu là trắng
                pages_to_remove = [p["page_number"] for p in all_pages_info if p.get("is_blank", True)]
                
            # Đảm bảo pages_to_remove là list số nguyên và sắp xếp giảm dần (xóa từ cuối lên)
            try:
                pages_to_remove = [int(p) for p in pages_to_remove]
                pages_to_remove.sort(reverse=True)  # Sắp xếp giảm dần để xóa từ trang cuối lên
            except (ValueError, TypeError) as e:
                self.log(f"Lỗi chuyển đổi pages_to_remove thành số nguyên: {e}", error=True)
                return False
                
            if not pages_to_remove:
                self.log("Không có trang nào cần xóa.")
                return file_path
                
            self.log(f"Sẽ xóa {len(pages_to_remove)} trang: {pages_to_remove}")
            
            # Tạo tên file kết quả
            if not output_path:
                file_dir = os.path.dirname(file_path)
                file_name = os.path.basename(file_path)
                base_name, ext = os.path.splitext(file_name)
                output_path = os.path.join(file_dir, f"{base_name}_no_blank{ext}")
            
            # Lưu tệp tạm thời để xử lý bằng Word COM
            temp_file = file_path + ".temp.docx"
            self.document.save(temp_file)
            self.log(f"Đã lưu tệp tạm thời: {temp_file}")
            
            # Mở tệp Word bằng COM để xóa trang
            self.log("Đang khởi tạo Word COM Application để xóa trang")
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            try:
                self.log(f"Đang mở tệp bằng Word COM: {temp_file}")
                doc = word_app.Documents.Open(os.path.abspath(temp_file))
                
                # Ghi log thông tin về tổng số trang
                total_pages = doc.ComputeStatistics(2)  # 2 là wdStatisticPages
                self.log(f"Tổng số trang trong tài liệu: {total_pages}")
                
                # Kiểm tra xem các trang cần xóa có hợp lệ không
                valid_pages = [p for p in pages_to_remove if 1 <= p <= total_pages]
                if len(valid_pages) != len(pages_to_remove):
                    invalid_pages = [p for p in pages_to_remove if p not in valid_pages]
                    self.log(f"Phát hiện {len(invalid_pages)} trang không hợp lệ: {invalid_pages}", error=True)
                
                # Nếu không có trang hợp lệ, kết thúc
                if not valid_pages:
                    self.log("Không có trang hợp lệ để xóa.")
                    return file_path
                    
                # Xác định hằng số Word VBA
                wdGoToPage = 1
                wdDoNotSaveChanges = 0
                
                # Lấy đối tượng Selection
                selection = word_app.Selection
                
                # Xóa từng trang (từ cuối lên để tránh ảnh hưởng đến số trang)
                pages_deleted = 0
                
                # Thu thập thông tin về section của từng trang
                page_section_info = []
                for page_num in range(1, total_pages + 1):
                    selection.GoTo(What=wdGoToPage, Which=1, Count=page_num, Name="")
                    has_section_break = False
                    
                    # Lưu vị trí hiện tại
                    current_pos = selection.Start
                    
                    # Di chuyển đến cuối trang và kiểm tra
                    selection.EndKey(Unit=6)  # 6 là wdStory để đảm bảo di chuyển đến cuối tài liệu
                    if selection.Information(9):  # 9 là wdAtEndOfRowMarker - kiểm tra xem có phải ở cuối đoạn
                        selection.MoveEnd(Unit=1, Count=1)  # 1 là wdCharacter - mở rộng selection thêm 1 ký tự
                        if selection.Text.find("\f") >= 0 or selection.Text.find("\x0c") >= 0 or selection.Text.find(chr(12)) >= 0:
                            has_section_break = True
                    
                    # Quay lại vị trí ban đầu
                    selection.Start = current_pos
                    selection.End = current_pos
                    
                    # Lấy section hiện tại
                    current_section = selection.Information(3)  # 3 là wdActiveEndSectionNumber
                    
                    page_section_info.append({
                        "page_number": page_num,
                        "section_number": current_section,
                        "has_section_break": has_section_break
                    })
                
                for page_num in valid_pages:
                    try:
                        self.log(f"Đang xóa trang {page_num}...")
                        
                        # Kiểm tra có section break không
                        page_info = next((info for info in page_section_info if info["page_number"] == page_num), None)
                        has_section_break = page_info and page_info.get("has_section_break", False)
                        
                        if has_section_break:
                            self.log(f"Trang {page_num} chứa section break, sử dụng phương pháp đặc biệt")
                            
                            # Đối với trang có section break, chúng ta sẽ chuyển đổi section break thành continuous
                            selection.GoTo(What=wdGoToPage, Which=1, Count=page_num, Name="")
                            
                            # Nhảy đến cuối trang và tìm section break
                            selection.EndKey(Unit=9)  # 9 là wdSection - đi đến cuối section hiện tại
                            
                            # Lấy chỉ số section hiện tại
                            current_section = selection.Information(3)  # 3 là wdActiveEndSectionNumber
                            
                            if current_section < doc.Sections.Count:
                                # Đổi section break từ Next Page sang Continuous
                                doc.Sections(current_section + 1).PageSetup.SectionStart = 0  # 0 là wdSectionContinuous
                                self.log(f"Đã chuyển section break của trang {page_num} thành continuous")
                            else:
                                self.log(f"Trang {page_num} là section cuối cùng, không thể chuyển đổi section break")
                            
                            # Di chuyển tới trang sau (hoặc trang hiện tại nếu đây là trang cuối)
                            next_page = min(page_num + 1, total_pages)
                            selection.GoTo(What=wdGoToPage, Which=1, Count=next_page, Name="")
                            
                            # Cập nhật lại tổng số trang
                            total_pages = doc.ComputeStatistics(2)
                            self.log(f"Số trang mới sau khi xử lý section break: {total_pages}")
                            
                            # Cập nhật lại thông tin trang
                            page_section_info = []
                            for p in range(1, total_pages + 1):
                                selection.GoTo(What=wdGoToPage, Which=1, Count=p, Name="")
                                current_section = selection.Information(3)
                                page_section_info.append({
                                    "page_number": p,
                                    "section_number": current_section,
                                    "has_section_break": False  # Reset thông tin về section break
                                })
                                
                            # Đánh dấu đã xử lý trang này
                            pages_deleted += 1
                        else:
                            # Phương pháp xóa trang thông thường
                            # Di chuyển đến trang cần xóa
                            selection.GoTo(What=wdGoToPage, Which=1, Count=page_num, Name="")
                            
                            # Nếu đây là trang cuối, xóa từ vị trí cuối trang trước đến hết
                            if page_num == total_pages:
                                # Xác định vị trí cuối trang trước
                                if page_num > 1:
                                    selection.GoTo(What=wdGoToPage, Which=1, Count=page_num-1, Name="")
                                    # Di chuyển đến cuối trang hiện tại
                                    selection.EndKey(Unit=6)  # 6 là wdStory
                                    selection.TypeBackspace()  # Xóa dấu ngắt trang
                                else:
                                    # Nếu đây là trang duy nhất, xóa toàn bộ nội dung
                                    selection.HomeKey(Unit=6)  # 6 là wdStory
                                    selection.EndKey(Unit=6, Extend=1)  # Extend=1 là mở rộng selection
                                    selection.Delete()
                            else:
                                # Di chuyển đến đầu trang hiện tại
                                selection.HomeKey(Unit=9)  # 9 là wdSection
                                
                                # Tìm vị trí bắt đầu của trang tiếp theo
                                next_page = page_num + 1
                                selection.GoTo(What=wdGoToPage, Which=1, Count=next_page, Name="")
                                selection.HomeKey(Unit=9)  # 9 là wdSection
                                end_pos = selection.Start
                                
                                # Quay lại trang cần xóa
                                selection.GoTo(What=wdGoToPage, Which=1, Count=page_num, Name="")
                                selection.HomeKey(Unit=9)  # 9 là wdSection
                                start_pos = selection.Start
                                
                                # Chọn toàn bộ nội dung từ đầu trang hiện tại đến đầu trang tiếp theo
                                selection.Start = start_pos
                                selection.End = end_pos
                                
                                # Xóa nội dung đã chọn
                                selection.Delete()
                            
                            pages_deleted += 1
                        
                        self.log(f"Đã xóa trang {page_num} thành công.")
                        
                    except Exception as e:
                        self.log(f"Lỗi khi xóa trang {page_num}: {e}", error=True)
                        self.log(traceback.format_exc(), error=True)
                
                # Lưu tài liệu
                self.log(f"Đã xóa {pages_deleted}/{len(valid_pages)} trang.")
                
                # Lưu tài liệu đã sửa
                self.log(f"Đang lưu tài liệu đã sửa vào: {output_path}")
                doc.SaveAs(os.path.abspath(output_path))
                doc.Close(SaveChanges=False)
                
                # Xóa tệp tạm
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        self.log(f"Đã xóa tệp tạm: {temp_file}")
                except Exception as e:
                    self.log(f"Không thể xóa tệp tạm: {temp_file}. Lỗi: {e}", error=True)
                
                return output_path
                
            finally:
                # Đảm bảo đóng Word
                try:
                    word_app.Quit()
                    self.log("Đã đóng Word COM Application")
                except Exception as e:
                    self.log(f"Lỗi khi đóng Word: {e}", error=True)
                
        except Exception as e:
            error_msg = f"Lỗi khi xóa trang trắng: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            return False

    def _get_section_type_name(self, section_type):
        """Trả về tên loại section dựa vào giá trị WD_SECTION_START"""
        section_types = {
            WD_SECTION_START.CONTINUOUS: "Continuous",
            WD_SECTION_START.NEW_COLUMN: "New Column",
            WD_SECTION_START.NEW_PAGE: "New Page",
            WD_SECTION_START.EVEN_PAGE: "Even Page",
            WD_SECTION_START.ODD_PAGE: "Odd Page"
        }
        return section_types.get(section_type, f"Unknown ({section_type})")

    def _get_blank_page_reason(self, page_info):
        """Trả về lý do tại sao trang được xác định là trắng"""
        if not page_info or not page_info.get("is_blank", False):
            return "Không phải trang trắng"
            
        page_text = page_info.get("text_content", "")
        text_length = len(page_text.strip())
        has_section_break = page_info.get("contains_section_break", False)
        
        # Kiểm tra các lý do
        if has_section_break and text_length < 5:
            return "Chỉ chứa section break"
            
        if page_info.get("is_first_page_after_section_break", False):
            section_type = page_info.get("section_type", "Unknown")
            if section_type in ["Next Page", "Even Page", "Odd Page"]:
                return f"Trang đầu tiên sau section break kiểu {section_type}"
            
        if text_length == 0:
            return "Không có nội dung"
            
        if not any(c.isalnum() for c in page_text):
            return "Chỉ chứa khoảng trắng"
            
        if "\f" in page_text and text_length < 10:
            return "Chứa page break"
            
        # Kiểm tra các patterns biểu thị trang trắng
        common_blank_patterns = [
            "\f", # Form feed
            "\u000C", # Form feed unicode
            "\x0c", # Form feed hex
            "Page intentionally left blank",
            "Trang này được để trống",
            "Trang trống"
        ]
        
        if any(pattern in page_text for pattern in common_blank_patterns):
            return "Chứa văn bản biểu thị trang trắng"
            
        # Lọc ra các ký tự đặc biệt và khoảng trắng
        filtered_text = ''.join(c for c in page_text if c.isalnum())
        
        if len(filtered_text) <= 3 and text_length <= 10:
            return "Chỉ chứa số trang hoặc ký tự đơn lẻ"
            
        if len(filtered_text) < 5:
            return "Quá ít ký tự chữ/số"
            
        return "Không đủ nội dung có ý nghĩa" 

    def process_document(self, file_path, method="combined", output_path=None, pages_to_remove=None, sections_to_fix=None):
        """
        Xử lý tài liệu để loại bỏ trang trắng theo phương pháp được chỉ định.
        
        Args:
            file_path (str): Đường dẫn đến file Word
            method (str, optional): Phương pháp xử lý. Có thể là:
                - "section_fix": Sửa section break
                - "page_remove": Xóa trang trắng trực tiếp
                - "combined": Kết hợp cả hai phương pháp
            output_path (str, optional): Đường dẫn lưu file kết quả. Nếu None, tạo tên tự động.
            pages_to_remove (list, optional): Danh sách các trang cần xóa (chỉ dùng khi method="page_remove" hoặc "combined")
            sections_to_fix (list, optional): Danh sách các section cần sửa (chỉ dùng khi method="section_fix" hoặc "combined")
            
        Returns:
            dict: Kết quả xử lý
        """
        try:
            if not output_path:
                file_dir = os.path.dirname(file_path)
                file_name = os.path.basename(file_path)
                base_name, ext = os.path.splitext(file_name)
                output_path = os.path.join(file_dir, f"{base_name}_processed{ext}")
                
            # Mở và phân tích tài liệu
            self.open_document(file_path)
            sections_results = self.analyze_document()
            updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
            
            # Kiểm tra nếu không có trang trắng
            blank_pages = [p for p in all_pages_info if p.get("is_blank")]
            if not blank_pages:
                self.log("Không phát hiện trang trắng nào trong tài liệu")
                return {
                    "success": True,
                    "method": "none",
                    "message": "Không phát hiện trang trắng nào trong tài liệu",
                    "pages_removed": [],
                    "sections_fixed": [],
                    "blank_pages_count": 0,
                    "output_path": file_path
                }
                
            # Danh sách các trang trắng
            blank_page_numbers = [p["page_number"] for p in blank_pages]
            self.log(f"Đã phát hiện tổng cộng {len(blank_page_numbers)} trang trắng: {blank_page_numbers}")
            
            # Nếu pages_to_remove không được chỉ định, sử dụng tất cả trang trắng đã phát hiện
            if pages_to_remove is None:
                pages_to_remove = blank_page_numbers
                
            # Lọc để chỉ lấy các trang cần xóa mà có trong trang trắng
            pages_to_remove = [p for p in pages_to_remove if p in blank_page_numbers]
            
            # Xác định các section liên quan đến trang trắng
            problem_sections = set()
            for page in blank_pages:
                page_num = page.get("page_number")
                if page_num in pages_to_remove and page.get("section_number"):
                    problem_sections.add(page.get("section_number"))
                    
                    # Nếu trang có section break, thì section tiếp theo sẽ ảnh hưởng
                    if page.get("has_section_break", False) and page.get("section_number") < len(updated_sections):
                        problem_sections.add(page.get("section_number") + 1)
            
            problem_sections = list(problem_sections)
            problem_sections.sort()
            
            # Nếu sections_to_fix không được chỉ định, sử dụng tất cả section có vấn đề
            if sections_to_fix is None:
                sections_to_fix = problem_sections
                
            # Lọc để chỉ lấy các section cần sửa mà có trong section có vấn đề
            sections_to_fix = [s for s in sections_to_fix if s in problem_sections]
            
            # Tạo biến lưu trữ kết quả
            results = {
                "success": False,
                "method": method,
                "message": "",
                "pages_removed": [],
                "sections_fixed": [],
                "blank_pages_count": len(blank_page_numbers),
                "output_path": output_path
            }
            
            # Phân tích thêm các trang trắng để xác định nguyên nhân và phương pháp tốt nhất
            section_break_pages = []
            empty_content_pages = []
            special_format_pages = []
            
            for page in blank_pages:
                page_num = page.get("page_number")
                if page_num not in pages_to_remove:
                    continue
                    
                if page.get("has_section_break", False):
                    section_break_pages.append(page_num)
                elif page.get("only_has_whitespace", False) or page.get("contains_only_page_number", False):
                    empty_content_pages.append(page_num)
                else:
                    special_format_pages.append(page_num)
            
            self.log(f"Phân tích tài liệu: {len(blank_page_numbers)} trang trắng, {len(problem_sections)} section có thể gây vấn đề")
            self.log(f"Trang có section break: {section_break_pages}")
            self.log(f"Trang chỉ có khoảng trắng: {empty_content_pages}")
            self.log(f"Trang có định dạng đặc biệt: {special_format_pages}")
            
            # Xử lý theo phương pháp được chỉ định
            if method == "section_fix":
                self.log("Áp dụng phương pháp sửa section break")
                
                # Nếu không có section nào cần sửa
                if not sections_to_fix:
                    self.log("Không có section nào cần sửa.")
                    results["success"] = False
                    results["message"] = "Không có section nào cần sửa"
                    return results
                
                # Sửa các section break
                result_path = self.fix_sections(file_path, sections_to_fix, output_path)
                
                if result_path:
                    # Kiểm tra kết quả sau khi sửa
                    self.open_document(result_path)
                    updated_sections_after = self.analyze_document()
                    _, all_pages_after = self.detect_blank_pages(updated_sections_after)
                    
                    blank_pages_after = [p["page_number"] for p in all_pages_after if p.get("is_blank")]
                    pages_removed = [p for p in pages_to_remove if p not in blank_pages_after]
                    
                    results["success"] = len(pages_removed) > 0
                    results["sections_fixed"] = sections_to_fix
                    results["pages_removed"] = pages_removed
                    results["output_path"] = result_path
                    results["message"] = f"Đã sửa {len(sections_to_fix)} section và loại bỏ {len(pages_removed)}/{len(pages_to_remove)} trang trắng"
                else:
                    results["success"] = False
                    results["message"] = "Lỗi khi sửa các section break"
                    
            elif method == "page_remove":
                self.log("Áp dụng phương pháp xóa trang trắng trực tiếp")
                
                # Nếu không có trang nào cần xóa
                if not pages_to_remove:
                    self.log("Không có trang nào cần xóa.")
                    results["success"] = False
                    results["message"] = "Không có trang nào cần xóa"
                    return results
                
                # Nếu có trang chứa section break, xử lý đặc biệt
                if section_break_pages:
                    self.log(f"Phát hiện {len(section_break_pages)} trang có section break, sử dụng xử lý đặc biệt")
                    
                    # Đầu tiên thử chuyển đổi section break
                    temp_path = file_path
                    if sections_to_fix:
                        temp_output = os.path.join(os.path.dirname(output_path), 
                                                   f"{os.path.splitext(os.path.basename(output_path))[0]}_temp{os.path.splitext(output_path)[1]}")
                        temp_path = self.fix_sections(file_path, sections_to_fix, temp_output)
                        if not temp_path:
                            temp_path = file_path
                    
                    # Sau đó xóa các trang
                    result_path = self.remove_blank_pages(temp_path, pages_to_remove, output_path)
                    
                    # Nếu có file tạm, xóa đi
                    if temp_path != file_path and os.path.exists(temp_path):
                        try:
                            os.remove(temp_path)
                            self.log(f"Đã xóa file tạm: {temp_path}")
                        except Exception as e:
                            self.log(f"Không thể xóa file tạm: {temp_path}. Lỗi: {e}", error=True)
                else:
                    # Xóa trang trắng thông thường
                    result_path = self.remove_blank_pages(file_path, pages_to_remove, output_path)
                
                if result_path:
                    # Kiểm tra kết quả sau khi xóa
                    self.open_document(result_path)
                    updated_sections_after = self.analyze_document()
                    _, all_pages_after = self.detect_blank_pages(updated_sections_after)
                    
                    blank_pages_after = [p["page_number"] for p in all_pages_after if p.get("is_blank")]
                    
                    # Lưu ý: số trang có thể thay đổi sau khi xóa, nên chúng ta kiểm tra tổng số trang trắng
                    original_blank_count = len([p for p in blank_pages if p["page_number"] in pages_to_remove])
                    remaining_blank_count = len(blank_pages_after)
                    pages_removed_count = original_blank_count - remaining_blank_count
                    
                    results["success"] = pages_removed_count > 0
                    results["pages_removed"] = [p for p in pages_to_remove if p not in blank_pages_after]
                    results["output_path"] = result_path
                    results["message"] = f"Đã xóa {pages_removed_count}/{len(pages_to_remove)} trang trắng"
                else:
                    results["success"] = False
                    results["message"] = "Lỗi khi xóa trang trắng"
                    
            elif method == "combined":
                self.log("Áp dụng phương pháp kết hợp: sửa section trước, sau đó xóa trang trắng còn lại")
                
                # Bước 1: Sửa section break
                if sections_to_fix:
                    self.log(f"Bước 1: Sửa {len(sections_to_fix)} section break")
                    temp_output = os.path.join(os.path.dirname(output_path), 
                                              f"{os.path.splitext(os.path.basename(output_path))[0]}_temp{os.path.splitext(output_path)[1]}")
                    result_path_step1 = self.fix_sections(file_path, sections_to_fix, temp_output)
                    
                    if not result_path_step1:
                        self.log("Lỗi khi sửa section break, chuyển sang phương pháp xóa trang trực tiếp")
                        result_path_step1 = file_path
                    else:
                        # Kiểm tra kết quả sau khi sửa section
                        self.open_document(result_path_step1)
                        updated_sections_after = self.analyze_document()
                        _, all_pages_after = self.detect_blank_pages(updated_sections_after)
                        
                        blank_pages_after = [p for p in all_pages_after if p.get("is_blank")]
                        blank_page_numbers_after = [p["page_number"] for p in blank_pages_after]
                        
                        pages_fixed_by_section = [p for p in pages_to_remove if p not in blank_page_numbers_after]
                        
                        if pages_fixed_by_section:
                            self.log(f"Sửa section đã giải quyết {len(pages_fixed_by_section)} trang: {pages_fixed_by_section}")
                            
                        # Cập nhật lại danh sách trang cần xóa
                        pages_to_remove = [p for p in pages_to_remove if p in blank_page_numbers_after]
                else:
                    self.log("Không có section nào cần sửa, chuyển sang phương pháp xóa trang trực tiếp")
                    result_path_step1 = file_path
                
                # Bước 2: Xóa trang trắng còn lại
                if pages_to_remove:
                    self.log(f"Bước 2: Xóa {len(pages_to_remove)} trang trắng còn lại")
                    self.log(f"Các trang sẽ xóa: {pages_to_remove}")
                    result_path_step2 = self.remove_blank_pages(result_path_step1, pages_to_remove, output_path)
                    
                    # Nếu có file tạm, xóa đi
                    if result_path_step1 != file_path and result_path_step1 != output_path and os.path.exists(result_path_step1):
                        try:
                            os.remove(result_path_step1)
                            self.log(f"Đã xóa file tạm: {result_path_step1}")
                        except Exception as e:
                            self.log(f"Không thể xóa file tạm: {result_path_step1}. Lỗi: {e}", error=True)
                            
                    if result_path_step2:
                        # Kiểm tra kết quả sau khi xóa
                        self.open_document(result_path_step2)
                        updated_sections_after = self.analyze_document()
                        _, all_pages_after = self.detect_blank_pages(updated_sections_after)
                        
                        blank_pages_after = [p["page_number"] for p in all_pages_after if p.get("is_blank")]
                        
                        # Tính toán trang đã xóa, nhưng lưu ý số trang có thể thay đổi
                        pages_removed = [p for p in blank_page_numbers if p not in blank_pages_after]
                        pages_remaining = [p for p in blank_page_numbers if p in blank_pages_after]
                        
                        results["success"] = len(pages_removed) > 0
                        results["sections_fixed"] = sections_to_fix
                        results["pages_removed"] = pages_removed
                        results["output_path"] = result_path_step2
                        results["remaining_pages"] = pages_remaining
                        results["message"] = f"Đã sửa {len(sections_to_fix)} section và xóa {len(pages_removed)}/{len(blank_page_numbers)} trang trắng"
                    else:
                        results["success"] = False
                        results["message"] = "Lỗi khi xóa trang trắng"
                else:
                    # Nếu không còn trang nào để xóa sau khi sửa section
                    results["success"] = True
                    results["sections_fixed"] = sections_to_fix
                    results["pages_removed"] = [p for p in blank_page_numbers if p not in pages_to_remove]
                    results["output_path"] = result_path_step1
                    results["message"] = f"Đã sửa {len(sections_to_fix)} section và loại bỏ tất cả {len(results['pages_removed'])} trang trắng"
            else:
                results["success"] = False
                results["message"] = f"Phương pháp không hợp lệ: {method}"
                
            return results
            
        except Exception as e:
            error_msg = f"Lỗi khi xử lý tài liệu: {e}"
            self.log(error_msg, error=True)
            self.log(traceback.format_exc(), error=True)
            
            return {
                "success": False,
                "method": method,
                "message": error_msg,
                "pages_removed": [],
                "sections_fixed": [],
                "blank_pages_count": len(blank_page_numbers) if 'blank_page_numbers' in locals() else 0,
                "output_path": None
            }

    def comprehensive_process(self, file_path, marked_pages=None, output_path=None, max_attempts=5):
        """
        Xử lý triệt để trang trắng bằng mọi phương pháp có thể để đảm bảo tất cả trang được đánh dấu đều được xử lý
        
        Args:
            file_path (str): Đường dẫn đến file Word
            marked_pages (list): Danh sách các trang đã được đánh dấu đặc biệt cần được xử lý bằng mọi giá
            output_path (str, optional): Đường dẫn lưu file kết quả. Nếu None, tạo tên tự động.
            max_attempts (int): Số lần thử tối đa, mặc định là 5
            
        Returns:
            dict: Kết quả xử lý với các thông tin:
                - success (bool): True nếu xử lý thành công
                - path (str): Đường dẫn đến file kết quả
                - processed_pages (list): Danh sách các trang đã xử lý thành công
                - remaining_pages (list): Danh sách các trang không thể xử lý
                - attempts (int): Số lần thử đã thực hiện
        """
        self.log(f"Bắt đầu xử lý triệt để với {len(marked_pages) if marked_pages else 'tất cả'} trang đánh dấu đặc biệt")
        
        # Nếu không cung cấp marked_pages, xem như không có trang nào cần xử lý đặc biệt
        if not marked_pages:
            self.log("Không có trang nào được đánh dấu đặc biệt để xử lý")
            result = {
                "success": False,
                "path": file_path,
                "processed_pages": [],
                "remaining_pages": [],
                "attempts": 0,
                "message": "Không có trang nào được đánh dấu đặc biệt để xử lý"
            }
            return result
            
        # Tạo tên file kết quả nếu chưa chỉ định
        if not output_path:
            file_dir = os.path.dirname(file_path)
            file_name = os.path.basename(file_path)
            base_name, ext = os.path.splitext(file_name)
            final_output_path = os.path.join(file_dir, f"{base_name}_processed_complete{ext}")
        else:
            final_output_path = output_path
            
        # Đảm bảo marked_pages là danh sách số nguyên và được sắp xếp
        try:
            marked_pages = [int(p) for p in marked_pages]
            marked_pages.sort()  # Sắp xếp tăng dần
        except (ValueError, TypeError) as e:
            self.log(f"Lỗi khi chuyển đổi marked_pages thành số nguyên: {e}", error=True)
            result = {
                "success": False,
                "path": file_path,
                "processed_pages": [],
                "remaining_pages": marked_pages,
                "attempts": 0,
                "message": f"Lỗi dữ liệu đầu vào: {e}"
            }
            return result
            
        # Danh sách các phương pháp sẽ thử, theo thứ tự từ ít xâm lấn đến xâm lấn nhất
        methods = ["section_fix", "combined", "page_remove"]
        
        # Biến theo dõi
        current_file = file_path
        processed_pages = []
        remaining_pages = marked_pages.copy()
        temp_files = []  # Danh sách các file tạm
        
        # Biến kiểm tra hiệu quả
        total_attempts = 0
        made_progress = True
        
        # Xử lý lặp lại cho đến khi không còn trang nào hoặc đạt đến giới hạn số lần thử
        while made_progress and remaining_pages and total_attempts < max_attempts:
            total_attempts += 1
            made_progress = False
            
            self.log(f"Lần thử {total_attempts}/{max_attempts}: Còn {len(remaining_pages)} trang cần xử lý")
            self.log(f"Các trang còn lại: {remaining_pages}")
            
            # Thử từng phương pháp
            for method_idx, method in enumerate(methods):
                if not remaining_pages:
                    break
                    
                self.log(f"Lần {total_attempts}, bước {method_idx+1}/{len(methods)}: Thử phương pháp {method}")
                
                # Tạo tên file tạm thời cho bước này
                file_dir = os.path.dirname(current_file)
                file_name = os.path.basename(current_file)
                base_name, ext = os.path.splitext(file_name)
                interim_path = os.path.join(file_dir, f"{base_name}_lần{total_attempts}_bước{method_idx+1}{ext}")
                temp_files.append(interim_path)  # Thêm file tạm vào danh sách
                
                # Mở và phân tích tài liệu hiện tại
                self.open_document(current_file)
                sections_results = self.analyze_document()
                updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
                
                # Xác định các section liên quan đến trang cần xử lý
                problem_sections = set()
                for page_num in remaining_pages:
                    # Tìm trang trong all_pages_info
                    for page in all_pages_info:
                        if page.get("page_number") == page_num and page.get("section_number"):
                            section_num = page.get("section_number")
                            problem_sections.add(section_num)
                            # Thêm cả section tiếp theo nếu có
                            if section_num < len(updated_sections):
                                problem_sections.add(section_num + 1)
                
                # Chuyển sang list và sắp xếp
                problem_sections = list(problem_sections)
                problem_sections.sort()
                
                # Thử áp dụng phương pháp hiện tại
                result = self.process_document(
                    current_file,
                    method=method,
                    output_path=interim_path,
                    pages_to_remove=remaining_pages,
                    sections_to_fix=problem_sections
                )
                
                if result and result != current_file:
                    # Cập nhật đường dẫn file hiện tại
                    current_file = result
                    
                    # Kiểm tra xem trang nào đã được xử lý
                    self.open_document(current_file)
                    sections_results = self.analyze_document()
                    updated_sections, new_pages_info = self.detect_blank_pages(sections_results)
                    
                    # Lấy số trang của tất cả các trang trắng còn lại
                    remaining_blank_pages = [p.get("page_number") for p in new_pages_info if p.get("is_blank", False)]
                    
                    # Xác định các trang đã được xử lý trong bước này
                    newly_processed = []
                    still_remaining = []
                    
                    for page_num in remaining_pages:
                        if page_num not in remaining_blank_pages:
                            newly_processed.append(page_num)
                        else:
                            still_remaining.append(page_num)
                    
                    # Cập nhật các danh sách
                    if newly_processed:
                        made_progress = True
                        processed_pages.extend(newly_processed)
                        remaining_pages = still_remaining
                        
                        self.log(f"Phương pháp {method} đã xử lý được {len(newly_processed)} trang: {newly_processed}")
                        self.log(f"Còn lại {len(remaining_pages)} trang: {remaining_pages}")
                
                # Nếu đã xử lý hết tất cả các trang, dừng vòng lặp
                if not remaining_pages:
                    self.log("Đã xử lý hết tất cả các trang đánh dấu đặc biệt!")
                    break
            
            # Nếu không có tiến triển nào trong lần lặp này, dừng vòng lặp
            if not made_progress:
                self.log(f"Không có tiến triển nào trong lần thử {total_attempts}, dừng xử lý")
                break
        
        # Đổi tên file kết quả cuối cùng
        if current_file != file_path:
            try:
                if os.path.exists(final_output_path):
                    os.remove(final_output_path)
                    
                # Kiểm tra xem đường dẫn có khác không
                if final_output_path != current_file:
                    os.rename(current_file, final_output_path)
                    self.log(f"Đã đổi tên file kết quả từ {current_file} thành {final_output_path}")
                    current_file = final_output_path
            except Exception as e:
                self.log(f"Không thể đổi tên file kết quả: {e}", error=True)
        
        # Xóa tất cả các file tạm
        self.log(f"Bắt đầu xóa {len(temp_files)} file tạm thời...")
        for temp_file in temp_files:
            if temp_file != current_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    self.log(f"Đã xóa file tạm: {temp_file}")
                except Exception as e:
                    self.log(f"Không thể xóa file tạm {temp_file}: {e}", error=True)
        
        # Tạo kết quả trả về
        success = len(remaining_pages) == 0
        message = ""
        
        if success:
            message = f"Đã xử lý thành công tất cả {len(processed_pages)} trang đánh dấu đặc biệt sau {total_attempts} lần thử"
        else:
            message = f"Đã xử lý {len(processed_pages)}/{len(processed_pages) + len(remaining_pages)} trang đánh dấu đặc biệt. Còn {len(remaining_pages)} trang không thể xử lý sau {total_attempts} lần thử"
        
        self.log(message)
        
        result = {
            "success": success,
            "path": current_file,
            "processed_pages": processed_pages,
            "remaining_pages": remaining_pages,
            "attempts": total_attempts,
            "message": message
        }
        
        return result 

    def advanced_process_with_tracking(self, file_path, marked_pages=None, output_path=None, callbacks=None):
        """
        Xử lý trang trắng với chức năng theo dõi và phân biệt giữa các trang đã xử lý và chưa xử lý.
        Phương thức này là phiên bản mở rộng của comprehensive_process với thêm khả năng cập nhật
        trạng thái theo thời gian thực.
        
        Args:
            file_path (str): Đường dẫn đến file Word
            marked_pages (list): Danh sách các trang đã được đánh dấu đặc biệt cần xử lý
            output_path (str, optional): Đường dẫn lưu file kết quả. Nếu None, tạo tên tự động.
            callbacks (dict, optional): Dictionary chứa các hàm callback:
                - on_page_processed(page_number): Gọi khi một trang đã được xử lý
                - on_page_failed(page_number): Gọi khi xử lý trang thất bại
                - on_progress(percent, message): Gọi để cập nhật tiến trình
                - on_complete(result): Gọi khi hoàn thành toàn bộ xử lý
            
        Returns:
            dict: Kết quả xử lý với các thông tin chi tiết như comprehensive_process
                cùng thêm thông tin chi tiết về từng trang:
                - page_details (list): Thông tin chi tiết về từng trang đã xử lý
        """
        self.log(f"Bắt đầu xử lý nâng cao với theo dõi cho {len(marked_pages) if marked_pages else 'tất cả'} trang")
        
        # Khởi tạo callbacks mặc định nếu không được cung cấp
        if callbacks is None:
            callbacks = {
                "on_page_processed": lambda page_number: None,
                "on_page_failed": lambda page_number: None,
                "on_progress": lambda percent, message: None,
                "on_complete": lambda result: None
            }
        else:
            # Đảm bảo tất cả callbacks đều tồn tại
            for callback_name in ["on_page_processed", "on_page_failed", "on_progress", "on_complete"]:
                if callback_name not in callbacks:
                    callbacks[callback_name] = lambda *args: None
        
        # Nếu không cung cấp marked_pages, xem như không có trang nào cần xử lý
        if not marked_pages:
            self.log("Không có trang nào được đánh dấu để xử lý")
            result = {
                "success": False,
                "path": file_path,
                "processed_pages": [],
                "remaining_pages": [],
                "attempts": 0,
                "message": "Không có trang nào được đánh dấu để xử lý",
                "page_details": []
            }
            callbacks["on_complete"](result)
            return result
            
        # Tạo tên file kết quả nếu chưa chỉ định
        if not output_path:
            file_dir = os.path.dirname(file_path)
            file_name = os.path.basename(file_path)
            base_name, ext = os.path.splitext(file_name)
            final_output_path = os.path.join(file_dir, f"{base_name}_processed_advanced{ext}")
        else:
            final_output_path = output_path
            
        # Đảm bảo marked_pages là danh sách số nguyên và được sắp xếp
        try:
            marked_pages = [int(p) for p in marked_pages]
            marked_pages.sort()  # Sắp xếp tăng dần
        except (ValueError, TypeError) as e:
            self.log(f"Lỗi khi chuyển đổi marked_pages thành số nguyên: {e}", error=True)
            result = {
                "success": False,
                "path": file_path,
                "processed_pages": [],
                "remaining_pages": marked_pages,
                "attempts": 0,
                "message": f"Lỗi dữ liệu đầu vào: {e}",
                "page_details": []
            }
            callbacks["on_complete"](result)
            return result
            
        # Khởi tạo thông tin chi tiết cho từng trang
        page_details = []
        for page_num in marked_pages:
            page_details.append({
                "page_number": page_num,
                "status": "pending",  # pending, processing, processed, failed
                "method_used": None,
                "attempts": 0,
                "last_error": None,
                "special_case": None  # Mới: Lưu thông tin nếu là trường hợp đặc biệt
            })
            
        # Thông báo bắt đầu xử lý
        callbacks["on_progress"](0, f"Bắt đầu xử lý {len(marked_pages)} trang...")
        
        # Danh sách các phương pháp sẽ thử, theo thứ tự từ ít xâm lấn đến xâm lấn nhất
        methods = ["section_fix", "combined", "page_remove", "special_cases"]
        
        # Biến theo dõi
        current_file = file_path
        processed_pages = []
        remaining_pages = marked_pages.copy()
        temp_files = []  # Danh sách các file tạm
        
        # Đánh dấu tất cả trang là đang chờ xử lý
        for page_info in page_details:
            callbacks["on_progress"](0, f"Đánh dấu trang {page_info['page_number']} cần xử lý...")
        
        # Biến kiểm tra hiệu quả
        total_attempts = 0
        max_attempts = 5  # Số lần thử tối đa
        made_progress = True
        
        # Xử lý lặp lại cho đến khi không còn trang nào hoặc đạt đến giới hạn số lần thử
        while made_progress and remaining_pages and total_attempts < max_attempts:
            total_attempts += 1
            made_progress = False
            
            self.log(f"Lần thử {total_attempts}/{max_attempts}: Còn {len(remaining_pages)} trang cần xử lý")
            callbacks["on_progress"](
                (total_attempts / max_attempts) * 100 * 0.5,  # Sử dụng nửa đầu của thanh tiến trình
                f"Lần thử {total_attempts}/{max_attempts}: Đang xử lý {len(remaining_pages)} trang còn lại..."
            )
            
            # Cập nhật trạng thái các trang còn lại là đang xử lý
            for page_num in remaining_pages:
                for page_info in page_details:
                    if page_info["page_number"] == page_num:
                        page_info["status"] = "processing"
                        page_info["attempts"] += 1
            
            # Thử từng phương pháp
            for method_idx, method in enumerate(methods):
                if not remaining_pages:
                    break
                    
                self.log(f"Lần {total_attempts}, bước {method_idx+1}/{len(methods)}: Thử phương pháp {method}")
                callbacks["on_progress"](
                    (total_attempts / max_attempts) * 100 * 0.5 + (method_idx / len(methods)) * 50 / max_attempts,
                    f"Lần {total_attempts}, bước {method_idx+1}/{len(methods)}: Thử phương pháp {method}..."
                )
                
                # Tạo tên file tạm thời cho bước này
                file_dir = os.path.dirname(current_file)
                file_name = os.path.basename(current_file)
                base_name, ext = os.path.splitext(file_name)
                interim_path = os.path.join(file_dir, f"{base_name}_lần{total_attempts}_bước{method_idx+1}{ext}")
                temp_files.append(interim_path)  # Thêm file tạm vào danh sách
                
                # Mở và phân tích tài liệu hiện tại
                self.open_document(current_file)
                sections_results = self.analyze_document()
                updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
                
                # Xác định các section liên quan đến trang cần xử lý
                problem_sections = set()
                for page_num in remaining_pages:
                    # Tìm trang trong all_pages_info
                    for page in all_pages_info:
                        if page.get("page_number") == page_num and page.get("section_number"):
                            section_num = page.get("section_number")
                            problem_sections.add(section_num)
                            # Thêm cả section tiếp theo nếu có
                            if section_num < len(updated_sections):
                                problem_sections.add(section_num + 1)
                
                # Chuyển sang list và sắp xếp
                problem_sections = list(problem_sections)
                problem_sections.sort()
                
                # Thử áp dụng phương pháp hiện tại
                try:
                    if method == "special_cases":
                        # Xử lý các trường hợp đặc biệt
                        self.log(f"Đang áp dụng xử lý đặc biệt cho {len(remaining_pages)} trang...")
                        
                        # Lấy thông tin chi tiết về các trang còn lại
                        special_pages = []
                        for page in all_pages_info:
                            if page.get("page_number") in remaining_pages:
                                # Phát hiện các đặc điểm đặc biệt
                                if (page.get("has_section_break") or 
                                    page.get("has_header_footer") or 
                                    page.get("has_watermark") or 
                                    page.get("is_protected")):
                                    special_pages.append(page.get("page_number"))
                        
                        if special_pages:
                            self.log(f"Đã phát hiện {len(special_pages)} trang có đặc điểm đặc biệt: {special_pages}")
                            
                            # Xử lý các trang đặc biệt
                            special_result = self.handle_special_blank_page_cases(
                                current_file, 
                                pages_to_process=special_pages,
                                output_path=interim_path
                            )
                            
                            if special_result["success"]:
                                self.log(f"Xử lý đặc biệt đã thành công cho {len(special_result['processed_pages'])} trang")
                                
                                # Cập nhật thông tin trang đã xử lý
                                for page_num in special_result["processed_pages"]:
                                    if page_num in remaining_pages:
                                        remaining_pages.remove(page_num)
                                        processed_pages.append(page_num)
                                        
                                        # Cập nhật trạng thái cho trang này
                                        for page_info in page_details:
                                            if page_info["page_number"] == page_num:
                                                page_info["status"] = "processed"
                                                page_info["method_used"] = "special_cases"
                                                page_info["special_case"] = special_result.get("special_case", "unknown")
                                                
                                        # Gọi callback cho trang đã xử lý
                                        callbacks["on_page_processed"](page_num)
                                        
                                # Cập nhật file hiện tại
                                if os.path.exists(interim_path):
                                    current_file = interim_path
                                    made_progress = True
                            else:
                                self.log("Xử lý đặc biệt không thành công")
                        else:
                            self.log("Không phát hiện trang đặc biệt, bỏ qua phương pháp này")
                    else:
                        result = self.process_document(
                            current_file,
                            method=method,
                            output_path=interim_path,
                            pages_to_remove=remaining_pages,
                            sections_to_fix=problem_sections
                        )
                        
                        if result["success"]:
                            self.log(f"Phương pháp {method} đã xử lý được {len(result['pages_removed'])} trang: {result['pages_removed']}")
                            
                            # Cập nhật trang đã xử lý
                            for page_num in result["pages_removed"]:
                                if page_num in remaining_pages:
                                    remaining_pages.remove(page_num)
                                    processed_pages.append(page_num)
                                    
                                    # Cập nhật trạng thái
                                    for page_info in page_details:
                                        if page_info["page_number"] == page_num:
                                            page_info["status"] = "processed"
                                            page_info["method_used"] = method
                                    
                                    # Gọi callback
                                    callbacks["on_page_processed"](page_num)
                            
                            # Cập nhật file hiện tại
                            if os.path.exists(result["output_path"]):
                                current_file = result["output_path"]
                                made_progress = True
                        else:
                            self.log(f"Phương pháp {method} không thành công: {result.get('message', 'Không có thông tin lỗi')}")
                    
                except Exception as e:
                    self.log(f"Lỗi khi áp dụng phương pháp {method}: {e}", error=True)
                    self.log(traceback.format_exc(), error=True)
                
                # Nếu đã xử lý tất cả trang, kết thúc
                if not remaining_pages:
                    break
            
            # Nếu không có tiến triển nào trong lần thử này, kết thúc
            if not made_progress:
                self.log(f"Không có tiến triển nào trong lần thử {total_attempts}, dừng xử lý")
                break
        
        # Đánh dấu các trang còn lại là thất bại
        for page_num in remaining_pages:
            # Cập nhật trạng thái
            for page_info in page_details:
                if page_info["page_number"] == page_num:
                    page_info["status"] = "failed"
                    page_info["last_error"] = "Không thể xử lý sau nhiều lần thử"
            
            # Gọi callback
            callbacks["on_page_failed"](page_num)
        
        # Đổi tên file kết quả cuối cùng
        if current_file != file_path:
            try:
                if os.path.exists(final_output_path):
                    os.remove(final_output_path)
                    
                # Kiểm tra xem đường dẫn có khác không
                if final_output_path != current_file:
                    os.rename(current_file, final_output_path)
                    self.log(f"Đã đổi tên file kết quả từ {current_file} thành {final_output_path}")
                    current_file = final_output_path
            except Exception as e:
                self.log(f"Không thể đổi tên file kết quả: {e}", error=True)
        
        # Xóa tất cả các file tạm
        self.log(f"Bắt đầu xóa {len(temp_files)} file tạm thời...")
        for temp_file in temp_files:
            if temp_file != current_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                    self.log(f"Đã xóa file tạm: {temp_file}")
                except Exception as e:
                    self.log(f"Không thể xóa file tạm {temp_file}: {e}", error=True)
        
        # Tạo kết quả trả về
        success = len(remaining_pages) == 0
        message = ""
        
        if success:
            message = f"Đã xử lý thành công tất cả {len(processed_pages)} trang đánh dấu đặc biệt sau {total_attempts} lần thử"
        else:
            message = f"Đã xử lý {len(processed_pages)}/{len(processed_pages) + len(remaining_pages)} trang đánh dấu đặc biệt. Còn {len(remaining_pages)} trang không thể xử lý sau {total_attempts} lần thử"
        
        self.log(message)
        
        result = {
            "success": success,
            "path": current_file,
            "processed_pages": processed_pages,
            "remaining_pages": remaining_pages,
            "attempts": total_attempts,
            "message": message,
            "page_details": page_details
        }
        
        # Gọi callback hoàn thành
        callbacks["on_complete"](result)
        
        return result

    def handle_special_blank_page_cases(self, file_path, pages_to_process, output_path=None):
        """
        Xử lý các trường hợp đặc biệt của trang trắng như trang có header/footer, 
        watermark, hoặc trang được bảo vệ.
        
        Args:
            file_path (str): Đường dẫn đến file Word
            pages_to_process (list): Danh sách các số trang cần xử lý
            output_path (str, optional): Đường dẫn lưu file kết quả
            
        Returns:
            dict: Kết quả xử lý
                - success (bool): True nếu xử lý thành công ít nhất một trang
                - processed_pages (list): Danh sách các trang đã xử lý thành công
                - failed_pages (list): Danh sách các trang không thể xử lý
                - output_path (str): Đường dẫn đến file kết quả
                - message (str): Thông báo kết quả
                - special_case (str): Loại trường hợp đặc biệt đã xử lý
        """
        if not pages_to_process:
            return {
                "success": False,
                "processed_pages": [],
                "failed_pages": [],
                "output_path": file_path,
                "message": "Không có trang nào để xử lý",
                "special_case": None
            }
            
        self.log(f"Bắt đầu xử lý đặc biệt cho {len(pages_to_process)} trang: {pages_to_process}")
        
        # Tạo đường dẫn đầu ra nếu không được cung cấp
        if not output_path:
            file_dir = os.path.dirname(file_path)
            file_name = os.path.basename(file_path)
            base_name, ext = os.path.splitext(file_name)
            output_path = os.path.join(file_dir, f"{base_name}_special_fixed{ext}")
        
        # Sao chép file ban đầu để xử lý
        try:
            shutil.copy2(file_path, output_path)
            self.log(f"Đã tạo bản sao của file tại {output_path}")
        except Exception as e:
            self.log(f"Lỗi khi tạo bản sao của file: {e}", error=True)
            return {
                "success": False,
                "processed_pages": [],
                "failed_pages": pages_to_process,
                "output_path": file_path,
                "message": f"Lỗi khi tạo bản sao: {e}",
                "special_case": None
            }
        
        # Mở tài liệu để phân tích
        self.open_document(file_path)
        sections_results = self.analyze_document()
        updated_sections, all_pages_info = self.detect_blank_pages(sections_results)
        
        # Tìm và phân loại các trang đặc biệt
        pages_with_header_footer = []
        pages_with_watermark = []
        pages_with_protection = []
        pages_with_special_breaks = []
        
        for page in all_pages_info:
            page_num = page.get("page_number")
            if page_num in pages_to_process:
                # Phân loại trang
                if page.get("has_header_footer", False):
                    pages_with_header_footer.append(page_num)
                if page.get("has_watermark", False):
                    pages_with_watermark.append(page_num)
                if page.get("is_protected", False):
                    pages_with_protection.append(page_num)
                if page.get("has_section_break", False) and page.get("break_type") != "Continuous":
                    pages_with_special_breaks.append(page_num)
        
        self.log(f"Phân loại trang đặc biệt:")
        self.log(f"- Trang có header/footer: {pages_with_header_footer}")
        self.log(f"- Trang có watermark: {pages_with_watermark}")
        self.log(f"- Trang được bảo vệ: {pages_with_protection}")
        self.log(f"- Trang có section break đặc biệt: {pages_with_special_breaks}")
        
        # Chọn phương pháp xử lý phù hợp dựa trên phân loại
        primary_special_case = None
        if pages_with_special_breaks:
            primary_special_case = "section_break"
        elif pages_with_header_footer:
            primary_special_case = "header_footer"
        elif pages_with_watermark:
            primary_special_case = "watermark"
        elif pages_with_protection:
            primary_special_case = "protection"
            
        # Chuẩn bị theo dõi kết quả xử lý
        processed_pages = []
        failed_pages = []
        made_progress = False
        
        # Lưu file tạm để xử lý với Word COM
        temp_file = os.path.join(os.path.dirname(output_path), f"temp_{os.path.basename(output_path)}")
        try:
            shutil.copy2(output_path, temp_file)
            self.log(f"Đã tạo file tạm tại {temp_file}")
        except Exception as e:
            self.log(f"Lỗi khi tạo file tạm: {e}", error=True)
            return {
                "success": False,
                "processed_pages": [],
                "failed_pages": pages_to_process,
                "output_path": output_path,
                "message": f"Lỗi khi tạo file tạm: {e}",
                "special_case": primary_special_case
            }
        
        try:
            # Khởi tạo Word COM
            self.log("Khởi tạo Word COM để xử lý trang đặc biệt")
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            # Mở tài liệu
            doc = word_app.Documents.Open(temp_file)
            self.log(f"Đã mở tài liệu, tổng số trang: {doc.ComputeStatistics(win32com.client.constants.wdStatisticPages)}")
            
            # Xử lý theo phương pháp chính được xác định
            if primary_special_case == "section_break":
                self.log("Áp dụng phương pháp xử lý section break đặc biệt")
                
                for page_num in pages_with_special_breaks:
                    try:
                        # Di chuyển đến trang cần xử lý
                        self.log(f"Di chuyển đến trang {page_num}")
                        word_app.Selection.GoTo(What=win32com.client.constants.wdGoToPage, Which=win32com.client.constants.wdGoToAbsolute, Count=page_num)
                        
                        # Tìm section break
                        selection = word_app.Selection
                        selection.Find.ClearFormatting()
                        selection.Find.Text = "^b"  # Section break
                        selection.Find.Forward = True
                        selection.Find.Wrap = win32com.client.constants.wdFindStop
                        
                        if selection.Find.Execute():
                            self.log(f"Trang {page_num}: Đã tìm thấy section break")
                            
                            # Xác định loại section break
                            current_section = selection.Sections[1]
                            break_type = current_section.PageSetup.SectionStart
                            
                            if break_type != win32com.client.constants.wdSectionContinuous:
                                self.log(f"Trang {page_num}: Đang chuyển section break từ loại {break_type} sang Continuous")
                                current_section.PageSetup.SectionStart = win32com.client.constants.wdSectionContinuous
                                made_progress = True
                                processed_pages.append(page_num)
                                self.log(f"Trang {page_num}: Đã chuyển đổi section break thành công")
                            else:
                                self.log(f"Trang {page_num}: Section break đã là Continuous, không cần thay đổi")
                                processed_pages.append(page_num)
                        else:
                            self.log(f"Trang {page_num}: Không tìm thấy section break dù được phát hiện trước đó")
                            failed_pages.append(page_num)
                    except Exception as e:
                        self.log(f"Lỗi khi xử lý section break trên trang {page_num}: {e}", error=True)
                        failed_pages.append(page_num)
            
            elif primary_special_case == "header_footer":
                self.log("Áp dụng phương pháp xử lý header/footer")
                
                for page_num in pages_with_header_footer:
                    try:
                        # Di chuyển đến trang cần xử lý
                        self.log(f"Di chuyển đến trang {page_num}")
                        word_app.Selection.GoTo(What=win32com.client.constants.wdGoToPage, Which=win32com.client.constants.wdGoToAbsolute, Count=page_num)
                        
                        # Xác định section hiện tại
                        selection = word_app.Selection
                        current_section = selection.Sections[1]
                        
                        # Xóa header và footer của section này
                        self.log(f"Trang {page_num}: Đang xóa header và footer")
                        
                        # Xóa header
                        for header_type in [win32com.client.constants.wdHeaderFooterPrimary, 
                                         win32com.client.constants.wdHeaderFooterFirstPage, 
                                         win32com.client.constants.wdHeaderFooterEvenPages]:
                            header = current_section.Headers(header_type)
                            if header.Exists:
                                header.Range.Delete()
                                self.log(f"Trang {page_num}: Đã xóa header loại {header_type}")
                        
                        # Xóa footer
                        for footer_type in [win32com.client.constants.wdHeaderFooterPrimary, 
                                         win32com.client.constants.wdHeaderFooterFirstPage, 
                                         win32com.client.constants.wdHeaderFooterEvenPages]:
                            footer = current_section.Footers(footer_type)
                            if footer.Exists:
                                footer.Range.Delete()
                                self.log(f"Trang {page_num}: Đã xóa footer loại {footer_type}")
                        
                        processed_pages.append(page_num)
                        made_progress = True
                        self.log(f"Trang {page_num}: Đã xóa header và footer thành công")
                    except Exception as e:
                        self.log(f"Lỗi khi xử lý header/footer trên trang {page_num}: {e}", error=True)
                        failed_pages.append(page_num)
            
            elif primary_special_case == "watermark":
                self.log("Áp dụng phương pháp xử lý watermark")
                
                for page_num in pages_with_watermark:
                    try:
                        # Di chuyển đến trang cần xử lý
                        self.log(f"Di chuyển đến trang {page_num}")
                        word_app.Selection.GoTo(What=win32com.client.constants.wdGoToPage, Which=win32com.client.constants.wdGoToAbsolute, Count=page_num)
                        
                        # Xóa tất cả watermark trong section hiện tại
                        selection = word_app.Selection
                        current_section = selection.Sections[1]
                        
                        # Xóa watermark (thường nằm trong header)
                        for header_type in [win32com.client.constants.wdHeaderFooterPrimary, 
                                       win32com.client.constants.wdHeaderFooterFirstPage, 
                                       win32com.client.constants.wdHeaderFooterEvenPages]:
                            try:
                                header = current_section.Headers(header_type)
                                if header.Exists:
                                    # Xóa tất cả hình ảnh trong header (thường là watermark)
                                    for shape in list(header.Shapes):
                                        shape.Delete()
                                        self.log(f"Trang {page_num}: Đã xóa shape/watermark từ header")
                            except Exception as inner_e:
                                self.log(f"Lỗi khi xóa watermark từ header trên trang {page_num}: {inner_e}")
                        
                        processed_pages.append(page_num)
                        made_progress = True
                        self.log(f"Trang {page_num}: Đã xóa watermark thành công")
                    except Exception as e:
                        self.log(f"Lỗi khi xử lý watermark trên trang {page_num}: {e}", error=True)
                        failed_pages.append(page_num)
            
            elif primary_special_case == "protection":
                self.log("Áp dụng phương pháp xử lý trang được bảo vệ")
                
                # Kiểm tra xem tài liệu có được bảo vệ không
                try:
                    is_protected = doc.ProtectionType != -1
                    if is_protected:
                        self.log("Tài liệu đang được bảo vệ. Đang cố gắng gỡ bỏ bảo vệ...")
                        try:
                            doc.Unprotect()
                            self.log("Đã gỡ bỏ bảo vệ tài liệu thành công")
                            is_protected = False
                        except Exception as e:
                            self.log(f"Không thể gỡ bỏ bảo vệ tài liệu: {e}", error=True)
                    
                    # Xử lý các trang được bảo vệ
                    if not is_protected:
                        for page_num in pages_with_protection:
                            try:
                                # Di chuyển đến trang cần xử lý
                                self.log(f"Di chuyển đến trang {page_num}")
                                word_app.Selection.GoTo(
                                    What=win32com.client.constants.wdGoToPage, 
                                    Which=win32com.client.constants.wdGoToAbsolute, 
                                    Count=page_num
                                )
                                
                                # Xóa nội dung trang
                                selection = word_app.Selection
                                selection.WholeStory()
                                selection.Delete()
                                
                                processed_pages.append(page_num)
                                made_progress = True
                                self.log(f"Trang {page_num}: Đã xóa nội dung trang được bảo vệ thành công")
                            except Exception as e:
                                self.log(f"Lỗi khi xử lý trang được bảo vệ {page_num}: {e}", error=True)
                                failed_pages.append(page_num)
                    else:
                        self.log("Không thể xử lý trang do tài liệu vẫn đang được bảo vệ")
                        failed_pages.extend(pages_with_protection)
                except Exception as e:
                    self.log(f"Lỗi khi xử lý trang được bảo vệ: {e}", error=True)
                    failed_pages.extend(pages_with_protection)
            
            # Lưu kết quả
            if processed_pages:
                self.log(f"Lưu tài liệu đã xử lý vào: {output_path}")
                doc.SaveAs(os.path.abspath(output_path))
                success = True
                message = f"Đã xử lý thành công {len(processed_pages)}/{len(pages_to_process)} trang trắng đặc biệt"
            else:
                success = False
                message = "Không có trang nào được xử lý thành công"
                output_path = file_path
            
            result = {
                "success": success,
                "message": message,
                "processed_pages": processed_pages,
                "failed_pages": failed_pages,
                "output_path": output_path,
                "special_case": primary_special_case
            }
            
            return result
        except Exception as e:
            self.log(f"Lỗi không xử lý được khi xử lý trang đặc biệt: {e}", error=True)
            return {
                "success": False,
                "message": f"Lỗi không xử lý được khi xử lý trang đặc biệt: {e}",
                "processed_pages": processed_pages,
                "failed_pages": failed_pages if failed_pages else pages_to_process,
                "output_path": file_path,
                "special_case": primary_special_case
            }
        finally:
            # Đóng tài liệu và Word
            try:
                doc.Close(SaveChanges=False)
            except:
                pass
                
            try:
                word_app.Quit()
            except:
                pass
            
            # Xóa tệp tạm
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    self.log(f"Đã xóa tệp tạm: {temp_file}")
            except Exception as e:
                self.log(f"Không thể xóa tệp tạm: {e}", error=True)
    
    def cleanup(self):
        """Dọn dẹp tài nguyên sau khi hoàn tất xử lý"""
        self.log("Đang dọn dẹp tài nguyên...")
        
        # Đóng tất cả các đối tượng Word và COM
        if hasattr(self, '_word_apps') and self._word_apps:
            for word_app in self._word_apps:
                try:
                    word_app.Quit()
                    self.log(f"Đã đóng Word Application")
                except Exception as e:
                    self.log(f"Lỗi khi đóng Word: {e}", error=True)
            
            # Xóa danh sách
            self._word_apps = []
        
        # Gọi garbage collector để giải phóng các đối tượng COM
        try:
            import gc
            gc.collect()
            self.log("Đã chạy garbage collector")
        except Exception as e:
            self.log(f"Lỗi khi chạy garbage collector: {e}", error=True)
        
        # Xóa file tạm
        if hasattr(self, 'temp_files') and self.temp_files:
            for temp_file in self.temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        self.log(f"Đã xóa file tạm: {temp_file}")
                except Exception as e:
                    self.log(f"Không thể xóa file tạm {temp_file}: {e}", error=True)
            
            # Xóa danh sách
            self.temp_files = []
            
        # Đảm bảo đối tượng document được đóng
        if hasattr(self, 'document'):
            self.document = None
            self.log("Đã xóa tham chiếu đến document")
        
        # Đóng đối tượng COM nếu cần
        try:
            import pythoncom
            pythoncom.CoUninitialize()
            self.log("Đã đóng COM")
        except Exception as e:
            self.log(f"Lỗi khi đóng COM: {e}", error=True)
            
        self.log("Hoàn tất dọn dẹp tài nguyên")
        
        # Đảm bảo rằng các file log được ghi đầy đủ
        try:
            # Force flush các log handler
            logging.shutdown()
        except Exception as e:
            print(f"Lỗi khi đóng logging: {e}")
            
        # Reset các tham chiếu để tránh leak memory
        try:
            self._initialized = False
            self._instance = None
        except:
            pass

    def save_temp_file(self):
        """Lưu file tạm để phân tích
        
        Returns:
            str: Đường dẫn đến file tạm, hoặc None nếu có lỗi
        """
        if not self.document or not self.file_path:
            self.log("Không thể lưu file tạm vì không có tài liệu", error=True)
            return None
            
        try:
            # Tạo tên file tạm độc nhất dựa trên timestamp
            import tempfile
            temp_dir = tempfile.gettempdir()
            temp_filename = f"autooffice_temp_{int(time.time())}.docx"
            temp_file = os.path.join(temp_dir, temp_filename)
            
            # Lưu tài liệu vào file tạm
            self.document.save(temp_file)
            self.log(f"Đã lưu file tạm: {temp_file}")
            
            # Đảm bảo danh sách temp_files tồn tại
            if not hasattr(self, 'temp_files'):
                self.temp_files = []
                
            # Thêm vào danh sách để xóa sau
            self.temp_files.append(temp_file)
            
            return temp_file
        except Exception as e:
            self.log(f"Lỗi khi lưu file tạm: {e}", error=True)
            self.log(traceback.format_exc(), error=True)
            return None