from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt
import os
import logging
from word_processor_2 import EmptyPageDetector, PageAnalyzer

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class WordProcessor:
    def __init__(self):
        self.document = None
        self.file_path = None
        self.sections_info = []
        self.empty_pages = []
        self.page_analyzer = None
        self.debug_mode = False
        
    def set_debug_mode(self, enabled=True):
        """Bật/tắt chế độ debug."""
        self.debug_mode = enabled
        if self.page_analyzer:
            self.page_analyzer.set_debug_mode(enabled)
        
    def open_document(self, file_path):
        """Mở tệp Word và đọc dữ liệu."""
        try:
            self.file_path = file_path
            self.document = Document(file_path)
            logger.info(f"Đã mở tệp: {file_path}")
            
            # Tạo phân tích trang
            self.page_analyzer = PageAnalyzer(file_path)
            # Áp dụng chế độ debug nếu có
            if self.debug_mode:
                self.page_analyzer.set_debug_mode(True)
            
            return True
        except Exception as e:
            logger.error(f"Lỗi khi mở tệp: {e}")
            return False
    
    def analyze_document(self):
        """Phân tích tài liệu để tìm các ngắt phần và trang trắng."""
        if not self.document:
            logger.error("Chưa mở tệp nào.")
            return False
            
        self.sections_info = []
        
        # Sử dụng công cụ phát hiện trang trắng nâng cao
        try:
            if self.page_analyzer:
                analysis_result = self.page_analyzer.analyze()
                self.empty_pages = analysis_result['empty_pages']
                
                # Log cấu trúc tài liệu để debug
                document_structure = analysis_result['document_structure']
                
                # Chỉ log nếu ở chế độ debug
                if self.debug_mode:
                    logger.info(f"Cấu trúc tài liệu:\n{document_structure}")
                else:
                    logger.info(f"Đã hoàn thành phân tích cấu trúc tài liệu")
        except Exception as e:
            logger.error(f"Lỗi khi phân tích trang trắng: {e}")
            import traceback
            logger.error(traceback.format_exc())
        
        # Lấy thông tin về các section trong tài liệu
        for i, section in enumerate(self.document.sections):
            section_type = section.start_type
            
            # Kiểm tra xem phần này có trong danh sách trang trắng không
            is_empty_page = any(page['section_index'] == i for page in self.empty_pages)
            
            # Nếu phát hiện là trang trắng hoặc là ngắt phần Next Page,
            # đánh dấu để hiển thị trong giao diện
            needs_conversion = is_empty_page
            if section_type == WD_SECTION_START.NEW_PAGE and not is_empty_page:
                # Đánh dấu là ngắt phần thông thường (không phải trang trắng)
                pass
            
            section_info = {
                "index": i,
                "type": section_type,
                "type_name": self._get_section_type_name(section_type),
                "page_break": section_type == WD_SECTION_START.NEW_PAGE,
                "needs_conversion": needs_conversion,
                "is_empty_page": is_empty_page
            }
            self.sections_info.append(section_info)
            
        logger.info(f"Đã phân tích tệp: Tìm thấy {len(self.sections_info)} phần.")
        
        # Đếm số trang trắng được phát hiện
        empty_pages_count = len(self.empty_pages)
        logger.info(f"Phát hiện {empty_pages_count} trang trắng trong tài liệu.")
        
        # Hiển thị thông tin chi tiết về mỗi trang trắng
        if empty_pages_count > 0:
            for i, page in enumerate(self.empty_pages):
                logger.info(f"Trang trắng {i+1}: Phần {page['section_index']+1}, "
                          f"Phương pháp phát hiện: {page.get('detection_method', 'unknown')}")
        
        return self.sections_info
    
    def _get_section_type_name(self, section_type):
        """Trả về tên kiểu ngắt phần."""
        section_types = {
            WD_SECTION_START.CONTINUOUS: "Continuous",
            WD_SECTION_START.NEW_COLUMN: "New Column",
            WD_SECTION_START.NEW_PAGE: "Next Page",
            WD_SECTION_START.EVEN_PAGE: "Even Page",
            WD_SECTION_START.ODD_PAGE: "Odd Page"
        }
        return section_types.get(section_type, "Unknown")
    
    def fix_empty_pages(self):
        """Sửa các trang trắng bằng cách chuyển ngắt phần sang Continuous."""
        if not self.document:
            logger.error("Chưa mở tệp nào.")
            return False
            
        # Đảm bảo tài liệu đã được phân tích
        if not self.sections_info:
            self.analyze_document()
            
        # Nếu không có trang trắng được phát hiện
        if not self.empty_pages and self.page_analyzer:
            logger.info("Không tìm thấy trang trắng trong phân tích trước đó, thực hiện phân tích lại...")
            analysis_result = self.page_analyzer.analyze()
            self.empty_pages = analysis_result['empty_pages']
            
            # Nếu vẫn không tìm thấy trang trắng
            if not self.empty_pages:
                logger.info("Không tìm thấy trang trắng để xử lý.")
                return 0
            
        # Sử dụng PageAnalyzer để sửa các trang trắng
        if self.page_analyzer and self.empty_pages:
            changes_made = self.page_analyzer.fix_empty_pages(self.document, self.empty_pages)
        else:
            # Phương pháp dự phòng nếu không có PageAnalyzer
            changes_made = 0
            # Chỉ chuyển đổi các ngắt phần gây ra trang trắng
            for i, section_info in enumerate(self.sections_info):
                if section_info["needs_conversion"] and section_info["is_empty_page"]:
                    section = self.document.sections[i]
                    section.start_type = WD_SECTION_START.CONTINUOUS
                    changes_made += 1
                    logger.info(f"Đã chuyển phần {i} từ 'Next Page' sang 'Continuous' (trang trắng)")
        
        # Cập nhật thông tin sections sau khi thay đổi
        self.update_sections_info_after_fix()
        
        logger.info(f"Đã thực hiện {changes_made} thay đổi để xóa trang trắng.")
        return changes_made
    
    def update_sections_info_after_fix(self):
        """Cập nhật thông tin các phần sau khi đã sửa."""
        if not self.document:
            return
            
        # Cập nhật thông tin về loại ngắt phần
        for i, section in enumerate(self.document.sections):
            if i < len(self.sections_info):
                self.sections_info[i]["type"] = section.start_type
                self.sections_info[i]["type_name"] = self._get_section_type_name(section.start_type)
                # Đánh dấu đã được sửa
                if section.start_type == WD_SECTION_START.CONTINUOUS and self.sections_info[i]["is_empty_page"]:
                    self.sections_info[i]["fixed"] = True
    
    def save_document(self, output_path=None):
        """Lưu tài liệu đã chỉnh sửa."""
        if not self.document:
            logger.error("Chưa mở tệp nào.")
            return False
            
        if not output_path:
            # Tạo tên tệp mới nếu không được chỉ định
            file_name, file_ext = os.path.splitext(self.file_path)
            output_path = f"{file_name}_fixed{file_ext}"
            
        try:
            self.document.save(output_path)
            logger.info(f"Đã lưu tệp vào: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Lỗi khi lưu tệp: {e}")
            return False
    
    def get_document_info(self):
        """Lấy thông tin cơ bản về tài liệu."""
        if not self.document:
            return None
            
        info = {
            "sections": len(self.document.sections),
            "paragraphs": len(self.document.paragraphs),
            "tables": len(self.document.tables),
            "empty_pages": len(self.empty_pages)
        }
        return info
