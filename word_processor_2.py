import os
import logging
import re
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2python import docx2python
import tempfile
import shutil
from docx.api import Document as ReadOnlyDocument
import win32com.client
import pythoncom
import sys
import comtypes.client
import time

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class EmptyPageDetector:
    """Class chuyên biệt để phát hiện trang trắng trong tài liệu Word."""
    
    def __init__(self):
        self.temp_dir = None
        self.debug_mode = False
        
    def set_debug_mode(self, enabled=True):
        """Bật/tắt chế độ debug để có thêm log."""
        self.debug_mode = enabled
        
    def extract_document(self, docx_path):
        """Trích xuất nội dung tài liệu Word để phân tích."""
        try:
            # Tạo thư mục tạm thời để giải nén
            self.temp_dir = tempfile.mkdtemp()
            logger.info(f"Tạo thư mục tạm thời: {self.temp_dir}")
            
            # Sử dụng docx2python để giải nén
            doc_data = docx2python(docx_path, self.temp_dir)
            
            return {
                'docx_data': doc_data,
                'temp_dir': self.temp_dir
            }
        except Exception as e:
            logger.error(f"Lỗi khi giải nén tài liệu: {e}")
            return None
    
    def cleanup(self):
        """Dọn dẹp các tệp tạm thời."""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
            logger.info(f"Đã xóa thư mục tạm thời: {self.temp_dir}")
    
    def get_page_count(self, docx_path):
        """Lấy số trang thực tế trong tài liệu Word."""
        try:
            # Phương pháp 1: Dùng COM để đếm số trang (chỉ hoạt động trên Windows với MS Office)
            try:
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(docx_path)
                page_count = doc.ComputeStatistics(2)  # 2 là wdStatisticPages
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
                logger.info(f"Số trang thực tế trong tài liệu: {page_count}")
                return page_count
            except Exception as e:
                logger.warning(f"Không thể đếm số trang bằng COM: {e}")
            
            # Phương pháp 2: Dùng comtypes
            try:
                # Yêu cầu Word tự động đóng
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(docx_path)
                page_count = doc.ComputeStatistics(2)  # wdStatisticPages
                doc.Close(False)
                word.Quit()
                logger.info(f"Số trang thực tế trong tài liệu (comtypes): {page_count}")
                return page_count
            except Exception as e:
                logger.warning(f"Không thể đếm số trang bằng comtypes: {e}")
                
            # Phương pháp thay thế: Ước lượng dựa trên số phần
            document = Document(docx_path)
            estimated_pages = len(document.sections)
            logger.info(f"Ước lượng số trang: {estimated_pages}")
            return estimated_pages
                
        except Exception as e:
            logger.error(f"Lỗi khi đếm số trang: {e}")
            return -1
            
    def detect_empty_pages_v2(self, docx_path):
        """Phương pháp cải tiến để phát hiện trang trắng chính xác hơn."""
        try:
            # Mở tài liệu
            document = Document(docx_path)
            
            # Thu thập thông tin cơ bản
            total_sections = len(document.sections)
            total_paragraphs = len(document.paragraphs)
            page_count = self.get_page_count(docx_path)
            
            logger.info(f"Tài liệu có {total_sections} phần, {total_paragraphs} đoạn văn, ước tính {page_count} trang")
            
            # Danh sách chứa phần nghi ngờ là trang trắng
            potential_empty_pages = []
            confirmed_empty_pages = []
            
            # Bước 1: Phân tích các ngắt phần với tiêu chí chặt chẽ
            for i, section in enumerate(document.sections):
                # Chỉ phân tích các phần kiểu Next Page
                if section.start_type == WD_SECTION_START.NEW_PAGE:
                    # Đánh dấu là "tiềm năng" để phân tích kỹ hơn
                    potential_empty_pages.append({
                        'section_index': i,
                        'type': section.start_type,
                    })
                    
            logger.info(f"Phát hiện {len(potential_empty_pages)} ngắt phần kiểu Next Page")
            
            # Bước 2: Phân tích sâu hơn các phần tiềm năng
            if potential_empty_pages:
                # Tạo danh sách các phần chứa nội dung thực
                section_has_content = self._analyze_section_content(document, docx_path)
                
                # Phân tích lại các phần tiềm năng
                for page in potential_empty_pages:
                    section_idx = page['section_index']
                    
                    # Nếu phân tích chỉ ra rằng phần này không có nội dung
                    if section_idx in section_has_content:
                        # Phần có nội dung, không phải trang trắng
                        if self.debug_mode:
                            logger.info(f"Phần {section_idx} có nội dung, không phải trang trắng")
                    else:
                        # Kiểm tra thêm nếu đây là phần đầu tiên hoặc cuối cùng
                        if section_idx == 0 or section_idx == total_sections - 1:
                            # Phần đầu/cuối thường không phải trang trắng
                            if self.debug_mode:
                                logger.info(f"Phần {section_idx} là phần đầu/cuối, có khả năng không phải trang trắng")
                            # Phân tích thêm
                            if self._is_definitely_empty(document, section_idx):
                                confirmed_empty_pages.append({
                                    'section_index': section_idx,
                                    'type': page['type'],
                                    'confidence': 'high',
                                    'detection_method': 'deep_analysis'
                                })
                        else:
                            # Phần giữa tài liệu, kiểm tra xem có phải trang trắng không
                            # Phần giữa có ngắt phần Next Page nhưng không có nội dung
                            if self._check_for_empty_middle_section(document, section_idx, section_has_content):
                                confirmed_empty_pages.append({
                                    'section_index': section_idx,
                                    'type': page['type'],
                                    'confidence': 'high',
                                    'detection_method': 'empty_middle_section'
                                })
            
            logger.info(f"Xác nhận {len(confirmed_empty_pages)} trang trắng sau khi phân tích kỹ lưỡng")
            
            # Hiển thị thông tin chi tiết về mỗi trang trắng được xác nhận
            for i, page in enumerate(confirmed_empty_pages):
                logger.info(f"Trang trắng {i+1}: Phần {page['section_index']+1}, " 
                          f"Phương pháp: {page.get('detection_method', 'unknown')}, "
                          f"Độ tin cậy: {page.get('confidence', 'medium')}")
                
            return confirmed_empty_pages
            
        except Exception as e:
            logger.error(f"Lỗi khi phát hiện trang trắng v2: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return []
    
    def _check_for_empty_middle_section(self, document, section_idx, section_has_content):
        """Kiểm tra xem một phần ở giữa tài liệu có phải là trang trắng không."""
        try:
            # Nếu phần không có trong danh sách phần có nội dung, kiểm tra thêm
            if section_idx not in section_has_content:
                # Kiểm tra phần trước và phần sau
                prev_has_content = (section_idx - 1) in section_has_content
                next_has_content = (section_idx + 1) in section_has_content
                
                # Nếu cả phần trước và phần sau đều có nội dung
                # nhưng phần này không có, có thể là trang trắng
                if prev_has_content and next_has_content:
                    logger.info(f"Phần {section_idx} nằm giữa hai phần có nội dung, có thể là trang trắng")
                    return True
                    
                # Kiểm tra thêm nếu phần trước hoặc phần sau là ngắt phần khác
                prev_section = document.sections[section_idx - 1] if section_idx > 0 else None
                next_section = document.sections[section_idx + 1] if section_idx < len(document.sections) - 1 else None
                
                # Nếu phần trước và phần này đều là ngắt phần Next Page
                if prev_section and prev_section.start_type == WD_SECTION_START.NEW_PAGE:
                    if self.debug_mode:
                        logger.info(f"Phần {section_idx} và phần trước đều là ngắt phần Next Page")
                    return True
                
                # Nếu phần có các đặc điểm đáng ngờ khác
                section = document.sections[section_idx]
                if self._is_definitely_empty(document, section_idx):
                    return True
                    
            return False
            
        except Exception as e:
            logger.error(f"Lỗi khi kiểm tra phần giữa {section_idx}: {e}")
            return False
    
    def _is_definitely_empty(self, document, section_idx):
        """Kiểm tra xem một phần có chắc chắn là trang trắng không."""
        try:
            # Lấy phần cần kiểm tra
            section = document.sections[section_idx]
            
            # Kiểm tra 1: Phần phải là ngắt phần Next Page
            if section.start_type != WD_SECTION_START.NEW_PAGE:
                return False
                
            # Kiểm tra 2: Không có header hoặc footer đặc biệt
            if section.different_first_page_header_footer:
                # Có header/footer trang đầu khác, không phải trang trắng
                return False
                
            # Kiểm tra 3: Không có các thuộc tính trang đặc biệt
            # (kiểm tra này có thể không hoàn hảo với python-docx)
            
            return True
            
        except Exception as e:
            logger.error(f"Lỗi khi kiểm tra trang trắng chắc chắn: {e}")
            return False
    
    def _analyze_section_content(self, document, docx_path):
        """Phân tích để xác định các phần có nội dung thực sự."""
        try:
            # Khởi tạo tập hợp phần có nội dung
            sections_with_content = set()
            
            # Phương pháp 1: Ước tính phần của mỗi đoạn văn có nội dung
            non_empty_paragraphs = []
            for i, para in enumerate(document.paragraphs):
                if para.text.strip():
                    non_empty_paragraphs.append(i)
            
            # Ước tính phân bố đoạn văn cho các phần
            if non_empty_paragraphs and len(document.sections) > 0:
                total_paras = len(document.paragraphs)
                total_sections = len(document.sections)
                
                # Phân bố đoạn văn theo phần dựa trên vị trí
                for para_idx in non_empty_paragraphs:
                    # Ước tính phần chứa đoạn văn này
                    est_section = min(int((para_idx / total_paras) * total_sections), total_sections - 1)
                    sections_with_content.add(est_section)
            
            # Phương pháp 2: Phân tích docx2python
            try:
                doc_data = docx2python(docx_path)
                for i, section in enumerate(doc_data.body_sections):
                    # Nếu phần có text, tables, hoặc hình ảnh
                    if self._contains_content(section):
                        sections_with_content.add(i)
            except Exception as e:
                logger.warning(f"Không thể phân tích nội dung bằng docx2python: {e}")
            
            logger.info(f"Tìm thấy {len(sections_with_content)} phần có nội dung")
            return sections_with_content
            
        except Exception as e:
            logger.error(f"Lỗi khi phân tích nội dung phần: {e}")
            return set()
    
    def _contains_content(self, section_content):
        """Kiểm tra xem phần có chứa nội dung không."""
        try:
            # Kiểm tra văn bản
            if isinstance(section_content, str) and section_content.strip():
                return True
                
            # Kiểm tra danh sách (đệ quy)
            if isinstance(section_content, list):
                for item in section_content:
                    if self._contains_content(item):
                        return True
                        
            # Kiểm tra bảng
            if hasattr(section_content, 'tables') and section_content.tables:
                for table in section_content.tables:
                    if table:  # Nếu bảng không trống
                        for row in table:
                            for cell in row:
                                if cell and cell.strip():
                                    return True
                                    
            return False
            
        except Exception as e:
            logger.error(f"Lỗi khi kiểm tra nội dung: {e}")
            return False
    
    def detect_empty_pages(self, docx_path):
        """Phát hiện trang trắng bằng nhiều phương pháp (phương pháp cũ)."""
        # Sử dụng phương pháp mới cải tiến
        return self.detect_empty_pages_v2(docx_path)
            
    def _analyze_section_for_emptiness(self, section, section_index, document, docx_data, non_empty_paragraphs):
        """Phân tích chi tiết một phần để kiểm tra tính trống rỗng."""
        try:
            # Kiểm tra thuộc tính của phần
            section_type = section.start_type
            
            # Nếu không phải là next page, không có khả năng tạo trang trắng
            if section_type != WD_SECTION_START.NEW_PAGE:
                return False
                
            # Kiểm tra xem phần có nội dung không (đoạn văn, bảng, hình ảnh)
            section_has_content = False
            
            # 1. Kiểm tra section body có trống không
            try:
                # Phương pháp với docx2python
                if section_index < len(docx_data.body_sections):
                    section_content = docx_data.body_sections[section_index]
                    for table in docx_data.body_sections[section_index].tables:
                        if table and any(cell.strip() for row in table for cell in row):
                            logger.info(f"Phần {section_index} có bảng với nội dung")
                            section_has_content = True
                            break
            except Exception as e:
                logger.warning(f"Không thể phân tích bảng trong phần {section_index}: {e}")
            
            # 2. Đếm số đoạn văn có nội dung trong phần
            # Đây là ước tính vì python-docx không cung cấp cách trực tiếp để xác định
            # đoạn văn thuộc phần nào
            section_paragraphs = 0
            if non_empty_paragraphs:
                # Ước tính dựa trên vị trí tương đối của đoạn văn
                paragraphs_per_section = len(non_empty_paragraphs) / len(document.sections)
                start_idx = int(section_index * paragraphs_per_section)
                end_idx = int((section_index + 1) * paragraphs_per_section)
                
                # Xem xét ngắt đoạn và ngắt trang trong nội dung
                for idx, content in non_empty_paragraphs[start_idx:end_idx]:
                    section_paragraphs += 1
            
            if section_paragraphs > 0:
                logger.info(f"Phần {section_index} có khoảng {section_paragraphs} đoạn văn")
                section_has_content = True
            
            # Nếu phần không có nội dung và là kiểu Next Page, đây có thể là trang trắng
            if not section_has_content:
                logger.info(f"Phần {section_index} có thể là trang trắng (không có nội dung)")
                return True
                
            return False
            
        except Exception as e:
            logger.error(f"Lỗi khi phân tích phần {section_index}: {e}")
            return False
    
    def _detect_empty_pages_from_xml(self, temp_dir, document, empty_pages_list):
        """Phát hiện trang trắng bằng cách phân tích cấu trúc XML."""
        try:
            # Kiểm tra file document.xml đã giải nén
            doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
            if not os.path.exists(doc_xml_path):
                logger.warning("Không tìm thấy file document.xml")
                return
                
            # Đọc nội dung XML
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()
                
            # Tìm các mẫu đặc trưng cho trang trắng
            # 1. Tìm các phần tử sectPr với thuộc tính type="nextPage"
            section_breaks = re.findall(r'<w:sectPr[^>]*>.*?<w:type\s+w:val="nextPage".*?</w:sectPr>', 
                                      xml_content, re.DOTALL)
            
            logger.info(f"Tìm thấy {len(section_breaks)} ngắt phần kiểu Next Page trong XML")
            
            # 2. Phân tích cấu trúc XML để tìm các phần không có nội dung
            # Tìm các phần tử <w:p> trống hoặc chỉ chứa định dạng
            empty_paragraphs = re.findall(r'<w:p>(\s*<w:pPr>.*?</w:pPr>\s*)?</w:p>', xml_content)
            logger.info(f"Tìm thấy {len(empty_paragraphs)} đoạn văn trống trong XML")
            
            # Tìm các mẫu XML đặc trưng của trang trắng
            empty_page_patterns = [
                # Mẫu trang trắng chứa ngắt phần tiếp theo là ngắt trang
                r'<w:p>.*?</w:p>\s*<w:sectPr>.*?<w:type\s+w:val="nextPage".*?</w:sectPr>\s*<w:p',
                # Mẫu trang trắng chỉ có header/footer
                r'<w:sectPr>.*?<w:headerReference.*?<w:footerReference.*?<w:type\s+w:val="nextPage".*?</w:sectPr>'
            ]
            
            for pattern in empty_page_patterns:
                matches = re.findall(pattern, xml_content, re.DOTALL)
                if matches:
                    logger.info(f"Tìm thấy {len(matches)} mẫu trang trắng có thể trong XML")
                    
                    # Thêm các trang trắng mới được phát hiện (giảm số lượng vì thường có false positives)
                    # Chỉ thêm 1-2 trang trắng từ phân tích XML, không thêm hết để tránh false positives
                    max_to_add = min(2, len(matches))
                    for i in range(max_to_add):
                        # Kiểm tra xem section_index đã được thêm vào chưa
                        section_indices = [ep['section_index'] for ep in empty_pages_list]
                        
                        # Ước tính section_index dựa trên vị trí trong XML
                        # Đây chỉ là ước tính và có thể không chính xác
                        est_section_index = min(i, len(document.sections) - 1)
                        
                        if est_section_index not in section_indices:
                            # Không tự động thêm từ XML, chỉ ghi log
                            if self.debug_mode:
                                logger.info(f"Phát hiện mẫu trang trắng có thể ở phần {est_section_index} từ XML")
            
        except Exception as e:
            logger.error(f"Lỗi khi phân tích XML: {e}")
    
    def _advanced_empty_page_detection(self, docx_path, empty_pages_list):
        """Phương pháp phát hiện trang trắng nâng cao."""
        try:
            document = Document(docx_path)
            
            # Phương pháp 1: Kiểm tra hình dạng của tài liệu
            # (các phần tiếp theo nhau với ngắt trang, không có bảng, hình ảnh)
            for i in range(len(document.sections) - 1):
                current_section = document.sections[i]
                next_section = document.sections[i + 1]
                
                # Kiểm tra nếu hai ngắt phần liên tiếp
                if (current_section.start_type == WD_SECTION_START.NEW_PAGE and 
                    next_section.start_type in [WD_SECTION_START.NEW_PAGE, 
                                               WD_SECTION_START.EVEN_PAGE, 
                                               WD_SECTION_START.ODD_PAGE]):
                    
                    # Kiểm tra thêm nếu phần hiện tại không có bảng, không có hình ảnh
                    # và không có nội dung đáng kể
                    
                    # Lưu ý: Đây là kiểm tra ước lượng vì python-docx không cho phép
                    # truy cập trực tiếp các phần tử của một phần cụ thể
                    is_potential_empty = True
                    
                    # Thêm vào danh sách nếu chưa có
                    section_indices = [ep['section_index'] for ep in empty_pages_list]
                    if i not in section_indices and is_potential_empty:
                        # Không tự động thêm, chỉ ghi log
                        if self.debug_mode:
                            logger.info(f"Phần {i} có thể là trang trắng (phương pháp kiểm tra phần liền kề)")
            
            # Phương pháp 2: Tìm kiếm các ngắt phần không có nội dung thông qua cấu trúc đoạn văn
            # Phương pháp này chỉ là ước lượng vì không thể liên kết trực tiếp đoạn văn với phần
            
            # Phương pháp 3: Phát hiện trang có khoảng trắng lớn
            # Tuy nhiên, điều này khó thực hiện chính xác với python-docx
            
            # Phương pháp heuristic: Kiểm tra nếu tổng số phần lớn hơn đáng kể so với số đoạn văn
            total_sections = len(document.sections)
            total_paragraphs = len(document.paragraphs)
            
            if total_sections > 3 and total_paragraphs / total_sections < 1.5:
                # Có thể có trang trắng, nhưng không tự động thêm tất cả
                # Chỉ log thông tin
                logger.info(f"Phát hiện tài liệu có thể có trang trắng (ratio phần/đoạn văn thấp: {total_paragraphs/total_sections:.2f})")
            
        except Exception as e:
            logger.error(f"Lỗi trong phát hiện trang trắng nâng cao: {e}")
    
    def visualize_document_structure(self, docx_path):
        """Tạo bản mô tả cấu trúc tài liệu để debug."""
        try:
            document = Document(docx_path)
            structure = []
            
            structure.append(f"=== Cấu trúc tài liệu ===")
            structure.append(f"Tổng số phần: {len(document.sections)}")
            structure.append(f"Tổng số đoạn văn: {len(document.paragraphs)}")
            structure.append(f"Tổng số bảng: {len(document.tables)}")
            structure.append(f"")
            
            for i, section in enumerate(document.sections):
                structure.append(f"--- Phần {i+1} ---")
                structure.append(f"Kiểu ngắt phần: {self._get_section_type_name(section.start_type)}")
                structure.append(f"Header khác nhau: {section.different_first_page_header_footer}")
                structure.append(f"Kích thước trang: {section.page_width.inches:.2f}\" x {section.page_height.inches:.2f}\"")
                structure.append(f"")
            
            structure.append(f"=== Phân bố nội dung ===")
            paragraph_count = 0
            for i, para in enumerate(document.paragraphs):
                if para.text.strip():
                    paragraph_count += 1
                    if paragraph_count <= 5 or paragraph_count > len(document.paragraphs) - 5:
                        structure.append(f"Đoạn văn {i+1}: '{para.text[:50]}...' (Dài: {len(para.text)})")
            
            if len(document.paragraphs) > 10:
                structure.append(f"... và {len(document.paragraphs) - 10} đoạn văn khác ...")
                
            # Thêm thông tin về phần có nội dung
            section_content = self._analyze_section_content(document, docx_path)
            structure.append(f"\n=== Phân bố nội dung theo phần ===")
            structure.append(f"Các phần có nội dung: {sorted(list(section_content))}")
            
            # Thêm thông tin về các trang trắng được phát hiện
            empty_pages = self.detect_empty_pages_v2(docx_path)
            structure.append(f"\n=== Trang trắng được phát hiện ===")
            for i, page in enumerate(empty_pages):
                structure.append(f"Trang trắng {i+1}: Phần {page['section_index']+1}, "
                               f"Phương pháp: {page.get('detection_method', 'unknown')}")
                
            return "\n".join(structure)
            
        except Exception as e:
            logger.error(f"Lỗi khi tạo cấu trúc tài liệu: {e}")
            return f"Không thể tạo cấu trúc tài liệu: {str(e)}"
    
    def _get_section_type_name(self, section_type):
        """Trả về tên kiểu ngắt phần."""
        section_types = {
            WD_SECTION_START.CONTINUOUS: "Continuous",
            WD_SECTION_START.NEW_COLUMN: "New Column",
            WD_SECTION_START.NEW_PAGE: "Next Page",
            WD_SECTION_START.EVEN_PAGE: "Even Page",
            WD_SECTION_START.ODD_PAGE: "Odd Page"
        }
        return section_types.get(section_type, "Unknown")
        
# Lớp mở rộng với công cụ phát hiện trang trắng tiên tiến
class PageAnalyzer:
    """Lớp phân tích trang trong tài liệu Word."""
    
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.empty_page_detector = EmptyPageDetector()
        
    def set_debug_mode(self, enabled=True):
        """Bật/tắt chế độ debug."""
        self.empty_page_detector.set_debug_mode(enabled)
        
    def analyze(self):
        """Phân tích toàn bộ tài liệu và trả về thông tin chi tiết."""
        empty_pages = self.empty_page_detector.detect_empty_pages(self.docx_path)
        document_structure = self.empty_page_detector.visualize_document_structure(self.docx_path)
        
        return {
            'empty_pages': empty_pages,
            'document_structure': document_structure
        }
        
    def fix_empty_pages(self, document, empty_pages):
        """Sửa các trang trắng được phát hiện."""
        changes_made = 0
        
        for page_info in empty_pages:
            section_index = page_info['section_index']
            
            # Kiểm tra giới hạn hợp lệ
            if 0 <= section_index < len(document.sections):
                section = document.sections[section_index]
                
                # Chỉ sửa các phần kiểu Next Page
                if section.start_type == WD_SECTION_START.NEW_PAGE:
                    section.start_type = WD_SECTION_START.CONTINUOUS
                    changes_made += 1
                    logger.info(f"Đã chuyển phần {section_index} từ 'Next Page' sang 'Continuous'")
        
        logger.info(f"Đã thực hiện {changes_made} thay đổi để xóa trang trắng.")
        return changes_made 